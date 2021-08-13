################################################################################
##
## BY: WANDERSON M.PIMENTA
## PROJECT MADE WITH: Qt Designer and PySide2
## V: 1.0.0
##
################################################################################

import os
import shutil
import sys
import platform
import re
import docx
from PySide2 import QtCore, QtGui, QtWidgets
from PySide2.QtCore import (QCoreApplication, QPropertyAnimation, QDate, QDateTime, QMetaObject, QObject, QPoint, QRect, QSize, QTime, QUrl, Qt, QEvent)
from PySide2.QtGui import (QBrush, QColor, QConicalGradient, QCursor, QFont, QFontDatabase, QIcon, QKeySequence, QLinearGradient, QPalette, QPainter, QPixmap, QRadialGradient)
from PySide2.QtWidgets import *
from openpyxl import load_workbook
import openpyxl
from datetime import date,timedelta,datetime
from decimal import Decimal
from openpyxl.styles import NamedStyle, Font, Border, Side, Alignment
import pandas as pd
import glob

# GUI FILE
from pandas import ExcelWriter

from ui_main import Ui_MainWindow

# IMPORT FUNCTIONS
from ui_functions import *
import time

class MainWindow(QMainWindow):
    def __init__(self):
        QMainWindow.__init__(self)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        # MOVE WINDOW
        def moveWindow(event):
            # RESTORE BEFORE MOVE
            if UIFunctions.returnStatus() == 1:
                UIFunctions.maximize_restore(self)

            # IF LEFT CLICK MOVE WINDOW
            if event.buttons() == Qt.LeftButton:
                self.move(self.pos() + event.globalPos() - self.dragPos)
                self.dragPos = event.globalPos()
                event.accept()
        def openInfoContracteAproapeDeTermen(event):
            if event.buttons() == Qt.LeftButton:
                self.goToContracteAproapeDeTermen()
                listaContracte =self.nrContractelorAproapeDeTermen(listaDB)
                if listaContracte[1] == "":
                    self.ui.label_textSusCercMijloc_10.hide()
                    self.ui.label_numarCercMijloc_7.setText("0")
                    self.ui.label_textJosCercMijloc_9.setText("Contracte expira azi")
                    self.ui.label_numarCercMijloc_7.setStyleSheet(u"border:none;\n"
                                                                   "color:rgb(220,220,220);\n"
                                                                   "padding-left:2px;\n"
                                                                   "padding-right:2px;\n"
                                                                   "font-size:70px")
                    self.ui.label_textJosCercMijloc_9.setStyleSheet(u"border:none;\n"
                                                                     "color:rgb(60,231,195);\n"
                                                                     "padding-bottom: 35px;")
                else:
                    self.ui.label_numarCercMijloc_7.setText(listaContracte[1])
                if listaContracte[2] == "":
                    self.ui.label_textSusCercMijloc_6.hide()
                    self.ui.label_numarCercMijloc_11.setText("0")
                    self.ui.label_textJosCercMijloc_5.setText("Contracte expira maine")
                    self.ui.label_numarCercMijloc_11.setStyleSheet(u"border:none;\n"
                                                                   "color:rgb(220,220,220);\n"
                                                                   "padding-left:2px;\n"
                                                                   "padding-right:2px;\n"
                                                                   "font-size:70px")
                    self.ui.label_textJosCercMijloc_5.setStyleSheet(u"border:none;\n"
                                                                     "color:rgb(60,231,195);\n"
                                                                     "padding-bottom: 35px;")
                else:
                    self.ui.label_numarCercMijloc_11.setText(listaContracte[2])
                if listaContracte[3] == "":
                    self.ui.label_textSusCercMijloc_11.hide()
                    self.ui.label_numarCercMijloc_12.setText("0")
                    self.ui.label_textJosCercMijloc_10.setText("Contracte expira peste 2 zile")
                    self.ui.label_numarCercMijloc_12.setStyleSheet(u"border:none;\n"
                                                                   "color:rgb(220,220,220);\n"
                                                                   "padding-left:2px;\n"
                                                                   "padding-right:2px;\n"
                                                                   "font-size:70px")
                    self.ui.label_textJosCercMijloc_10.setStyleSheet(u"border:none;\n"
                                                                     "color:rgb(60,231,195);\n"
                                                                     "padding-bottom: 35px;")
                else:
                    self.ui.label_numarCercMijloc_12.setText(listaContracte[3])
                if listaContracte[4] == "":
                    self.ui.label_textSusCercMijloc_14.hide()
                    self.ui.label_numarCercMijloc_13.setText("0")
                    self.ui.label_textJosCercMijloc_11.setText("Contracte expira peste 3 zile")
                    self.ui.label_numarCercMijloc_13.setStyleSheet(u"border:none;\n"
                                                                "color:rgb(220,220,220);\n"
                                                                "padding-left:2px;\n"
                                                                "padding-right:2px;\n"
                                                                "font-size:70px")
                    self.ui.label_textJosCercMijloc_11.setStyleSheet(u"border:none;\n"
                                                                  "color:rgb(60,231,195);\n"
                                                                  "padding-bottom: 35px;")
                else:
                    self.ui.label_numarCercMijloc_13.setText(listaContracte[4])
        def openGestiuneVanzari(event):
            if event.buttons() == Qt.LeftButton:
                self.goToGestiuneVanzari()
        def openInfoContracteInTermen(event):
            if event.buttons() == Qt.LeftButton:
                self.goToContracteInTermen()
        def openInfoProduseInTermen(event):
            if event.buttons() == Qt.LeftButton:
                self.goToProduseInTermen()
                self.ui.stackedWidget_5.setCurrentIndex(0)
                # self.produseInTermen(listaDB)
                # self.produseInAsteptare(listaDB)
        def selectCautaDupaCNP(event):
            if event.buttons() == Qt.LeftButton:
                self.ui.radioButton.setChecked(True)
                self.ui.radioButton_2.setChecked(False)
                self.ui.stackedWidget_2.setCurrentIndex(1)

        def selectCautaDupaNume(event):
            if event.buttons() == Qt.LeftButton:
                self.ui.radioButton.setChecked(False)
                self.ui.radioButton_2.setChecked(True)
                self.ui.stackedWidget_2.setCurrentIndex(0)

        def selectContractVanzare(event):
            if event.buttons() == Qt.LeftButton:
                self.ui.radioButton_3.setChecked(True)
                self.ui.radioButton_4.setChecked(False)
        def selectContractAmanet(event):
            if event.buttons() == Qt.LeftButton:
                self.ui.radioButton_3.setChecked(False)
                self.ui.radioButton_4.setChecked(True)


        # SET TITLE BAR
        self.ui.title_bar.mouseMoveEvent = moveWindow
        #START of  My code
    ######################################################################################
        # 1. Se fac copii la DB si Template Contracte in folderele specifice
        currentDir = os.getcwd()
        listaDB = "\\BazaDate\\DB.xlsx"
        copyLista = "\\BazaDate\\CopyDB"
        originalDB = currentDir + listaDB
        copyDB = currentDir + copyLista
        shutil.copy(originalDB, copyDB)
        listaDB = copyLista + "\\DB.xlsx"
        listaDB = currentDir + listaDB
        templateContract = "\\templateContract\\TemplateContractAmanet.docx"
        destinationFolder = "\\Contracte"
        originalTemplate = currentDir + templateContract
        copyTemplate = currentDir + destinationFolder
        shutil.copy(originalTemplate, copyTemplate)
    #####################################################################################
        # 2. Se executa functii necesare la initializarea aplicatiei
        # 2.1 Se verifica starea contractelor si a produselor in DB si se muta in sheet-urile specifice fiecare situatii
        self.verificareStareContractesiProduse(listaDB)

        # 2.2 Se verifica nr de contracte aproape de termen, nr de contracte in termen, si nr de produse in termen, pentru a afisa aceste numere in pagina principala
        self.nrContractelorAproapeDeTermen(listaDB)
        self.nrContracteInTermen(listaDB)
        self.produseInTermen(listaDB)

        # 2.3 Se populeaza tabelele din aplicatie cu valorile corespunzatoare
        self.PopulareTabelContracte(listaDB)
        self.PopulareTabelProduse(listaDB)
        self.populeazaTabelCautaClient(listaDB)
        self.PopulareTabelGestiuneVanzariProduseAmanet(listaDB)
        self.PopulareTabelGestiuneVanzareProduseFurnizor(listaDB)
        # Ajustez dimensiunea la tabelele din Gestiune vanzare
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(0, 75 )
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(1, 207 )
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(2, 70)
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(3, 126)
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(4, 106)
        self.ui.tableWidget_produseLichidateClient_6 .setColumnWidth(5, 102)
        self.ui.tableWidget_produseLichidateClient_3 .setColumnWidth(0, 231)
        self.ui.tableWidget_produseLichidateClient_3 .setColumnWidth(1, 80)
        self.ui.tableWidget_produseLichidateClient_3 .setColumnWidth(2, 40)
        self.ui.tableWidget_produseLichidateClient_3 .setColumnWidth(3, 40)
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(0,272)
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(1 ,80)
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(2 ,110 )
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(3 ,110)
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(4 ,98 )
        self.ui.tableWidget_produseLichidateClient_8.setColumnWidth(5 ,97 )
        # 2.4 Se seteaza ca initial sa se afiseze home page-ul si tabelele cu contracte/produse in termen
        self.ui.stackedWidget.setCurrentIndex(0)
        self.ui.stackedWidget_4.setCurrentIndex(0)
        self.ui.stackedWidget_5.setCurrentIndex(0)
        # 2.5 Se ataseaza actiunile necesare cand se face click pe cercurile din homePage
        self.ui.frame_cercStanga.mousePressEvent = openInfoContracteAproapeDeTermen
        self.ui.frame_cercDreapta.mousePressEvent = openInfoContracteInTermen
        self.ui.frame_cercMijloc.mousePressEvent = openInfoProduseInTermen
        self.ui.frame_GestiuneDeVanzari.mousePressEvent = openGestiuneVanzari

        # self.ui.label_numarCercDreapta.setText(str(self.ui.tableWidget_contracteInTermen.rowCount()))
        # Buttons from home page
        self.ui.pushButton_genereazaContract.clicked.connect(lambda : self.goToGenereazaContract(listaDB))
        # self.ui.pushButton_incheieContract.clicked.connect(lambda : self.goToIncheieContract())
        # self.ui.pushButton_prelungesteContract.clicked.connect(lambda : self.goToPrelungesteContract())
        self.ui.pushButton_cautaClient.clicked.connect(lambda : self.goToCautaClient())
        # Buttons from Genereaza contract
        self.ui.pushButton_backToHome.clicked.connect(lambda :self.goHomePage(listaDB))
        self.ui.pushButton_backToHome_2.clicked.connect(lambda: self.adaugaDateClientInContract(listaDB))
        self.ui.pushButton_adaugaProdus.clicked.connect(lambda :self.adaugaDateProdusInContract(listaDB))
        self.ui.pushButton_finalizareContract.clicked.connect(lambda :self.Finalizare(listaDB))
        self.ui.radioButton_3.mousePressEvent = selectContractVanzare
        self.ui.radioButton_4.mousePressEvent = selectContractAmanet
        # Buttons from Cauta client
        self.ui.pushButton_backFromClientGasit.clicked.connect(lambda : self.goHomePage(listaDB))
        self.ui.pushButton_finalizareContract_5.clicked.connect(lambda :self.clickCautaClient(listaDB))
        self.ui.pushButton_genereazaContractClientGasit.clicked.connect(lambda: self.goToGenereazaContract(listaDB))
        self.ui.radioButton.mousePressEvent = selectCautaDupaCNP
        self.ui.radioButton_2.mousePressEvent = selectCautaDupaNume
        # Buttons from Contracte In Termen
        self.ui.pushButton_maiMultiClientiGasitiOK_2.clicked.connect(lambda :self.goHomePage(listaDB))
        self.ui.pushButton_goHome_3.clicked.connect(lambda : self.goToTabelContracteInAsteptare())
        self.ui.pushButton_goHome_2.clicked.connect(lambda : self.goToTabelContractePrelungite())
        self.ui.pushButton_goHome_5.clicked.connect(lambda : self.goToTabelContracteExpirate())
        self.ui.pushButton_goHome_4.clicked.connect(lambda : self.goToTabelContracteInTermen())
        self.ui.pushButton_goHome_7.clicked.connect(lambda : self.goToTabelContracteIncheiate())
        self.ui.pushButton_goHome_6.clicked.connect(lambda : self.goToTabelContracteInAsteptare())
        self.ui.pushButton_goHome_9.clicked.connect(lambda : self.goToTabelContractePrelungite())
        self.ui.pushButton_goHome_8.clicked.connect(lambda : self.goToTabelContracteExpirate())
        self.ui.pushButton_goHome_16.clicked.connect(lambda : self.goToTabelContracteInTermen())
        self.ui.pushButton_goHome_15.clicked.connect(lambda: self.goToTabelContracteIncheiate())
        self.ui.pushButton_backToHome_5.clicked.connect(lambda : self.selecteazaContractInTabel())
        self.ui.pushButton_incheieContract_6.clicked.connect(lambda : self.prelungesteContract(listaDB))
        self.ui.pushButton_incheieContract_7.clicked.connect(lambda : self.incheieContract(listaDB))
        self.ui.pushButton_incheieContract_8.clicked.connect(lambda : self.finalizarePrelungireContract(listaDB))
        self.ui.pushButton_finalizareIncheiereContract.clicked.connect(lambda : self.finalizareIncheiereContract(listaDB))
        self.ui.pushButton_goHome.clicked.connect(lambda :self.goHomePage(listaDB))
        self.ui.pushButton_goHome_10.clicked.connect(lambda : self.calculeazaComisionDupaXZilePrelungire())
        self.ui.pushButton_lichideazaContractCuTermenDepasit.clicked.connect(lambda: self.lichideazaContractCuTermenDepasit(listaDB))
        # Buttons from Produse In Termen
        self.ui.pushButton_backToHome_6.clicked.connect(lambda: self.selecteazaContractInTabelProduse())
        self.ui.pushButton_maiMultiClientiGasitiOK_4.clicked.connect(lambda: self.goHomePage(listaDB))
        self.ui.pushButton_goHome_11.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(4))
        self.ui.pushButton_goHome_12.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(1))
        self.ui.pushButton_goHome_13.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(0))
        self.ui.pushButton_goHome_14.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(2))
        self.ui.pushButton_goHome_18.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(3))
        self.ui.pushButton_goHome_17.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(1))
        self.ui.pushButton_goHome_20.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(4))
        self.ui.pushButton_goHome_19.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(2))
        self.ui.pushButton_goHome_22.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(0))
        self.ui.pushButton_goHome_21.clicked.connect(lambda: self.ui.stackedWidget_5.setCurrentIndex(3))

        # Butoane din GESTIUNE DE VANZARI
        self.ui.pushButton_goHome_33.clicked.connect(lambda: self.goToPageProduseFurnizor())
        self.ui.pushButton_goHome_32.clicked.connect(lambda: self.goToPageProduseFurnizor())
        self.ui.pushButton_goHome_28.clicked.connect(lambda: self.goToPageProduseAmanet())
        self.ui.pushButton_goHome_27.clicked.connect(lambda: self.goToPageProduseAmanet())
        self.ui.pushButton_goHome_25.clicked.connect(lambda: self.goHomePage(listaDB))
        self.ui.pushButton_goHome_23.clicked.connect(lambda: self.goHomePage(listaDB))
        self.ui.pushButton_finalizareIncheiereContract_8.clicked.connect(lambda: self.adaugaProdusAmanetLaVanzare(listaDB))
        self.ui.pushButton_finalizareIncheiereContract_7.clicked.connect(lambda : self.goToPageVanzareProduseAmanet())
        self.ui.pushButton_finalizareIncheiereContract_10.clicked.connect(lambda : self.finalizareVinzareProdusAmanet(listaDB))
        self.ui.pushButton_goHome_26.clicked.connect(lambda : self.ui.stackedWidget_6.setCurrentIndex(1))
        self.ui.pushButton_goHome_31.clicked.connect(lambda : self.adaugaProdusFurnizor())
        self.ui.pushButton_finalizareIncheiereContract_5.clicked.connect(lambda : self.finalizareAdaugareProduseFurnizor(listaDB))
        self.ui.pushButton_finalizareIncheiereContract_2.clicked.connect(lambda : self.vindeProdusFurnizor(listaDB))
        # END of my code

        ## ==> SET UI DEFINITIONS
        UIFunctions.uiDefinitions(self)


        ## SHOW ==> MAIN WINDOW
        ########################################################################
        self.show()

    ## APP EVENTS
    ########################################################################
    def mousePressEvent(self, event):
        self.dragPos = event.globalPos()
    def goToGestiuneVanzari(self):
        self.ui.stackedWidget.setCurrentIndex(7)
        self.goToPageProduseAmanet()
    def goToPageProduseFurnizor(self):
        if not self.ui.tableWidget_produseLichidateClient_6.rowCount() == 0 and not self.ui.tableWidget_produseLichidateClient_6.rowCount() ==0:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Trebuie terminata procedura de adaugare a produselor in gestiunea de vanzari")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            self.ui.stackedWidget_6.setCurrentIndex(0)
    def goToPageProduseAmanet(self):
        self.ui.stackedWidget_6.setCurrentIndex(1)
    def verificareStareContractesiProduse(self,listaDB):
        # Functia ruleaza la initializarea aplicatiei pentru a updata starea contractelor ( daca dintre contractele in termen au trecut in asteptare/ din contractele in asteptare/prelungite au trecut in expirate)
        # 1. Pentru contractele in termen
        # listaContracte = ["ListaContracteInTermen", "ListaContractePrelungite", "ListaContracteInAsteptare"]
        currentDir = os.getcwd()
        lista = load_workbook(listaDB)
        now = date.today()
        df_listaContracte = pd.read_excel(listaDB, sheet_name="ListaContracteInTermen") # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        randuriDeSters = []
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            dataAsteptare = (dataExpirare + timedelta(days=5))
            if dataExpirare < now and dataAsteptare >= now: # Contractul este in asteptare si trebuie mutat in lista cu ContracteInAsteptare si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
                contract = df_listaContracte.loc[[j]]
                row =contract.values.tolist()
                ws = lista["ListaContracteInAsteptare"]
                ws.append(row[0])
                randuriDeSters.append(j+2)
            elif dataAsteptare < now: # Contractul a depasit si data asteptarii => este expirat
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaContracteExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaContracteInTermen"]
        for i in range(len(randuriDeSters)-1,-1,-1):
            ws.delete_rows(randuriDeSters[i])
        # 2. Pentru contractele in asteptare
        randuriDeSters = []
        df_listaContracte = pd.read_excel(listaDB,
                                          sheet_name="ListaContracteInAsteptare")  # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            dataAsteptare = (dataExpirare + timedelta(days=5))
            if dataAsteptare < now:  # Contractul este in asteptare si trebuie mutat in lista cu ContracteExpirate si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaContracteExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaContracteInAsteptare"]
        for i in range(len(randuriDeSters) - 1, -1, -1):
            ws.delete_rows(randuriDeSters[i])
        # 3. Pentru contractele prelungite
        randuriDeSters = []
        df_listaContracte = pd.read_excel(listaDB,
                                              sheet_name="ListaContractePrelungite")  # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            # dataAsteptare = (dataExpirare + timedelta(days=5))
            # if dataExpirare < now and dataAsteptare >= now:  # Contractul este in asteptare si trebuie mutat in lista cu ContracteInAsteptare si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
            #     contract = df_listaContracte.loc[[j]]
            #     row = contract.values.tolist()
            #     ws = lista["ListaContracteInAsteptare"]
            #     ws.append(row[0])
            #     randuriDeSters.append(j + 2)
            if dataExpirare < now:  # Contractul a fost prelungit => data expirarii a fost extinsa , iar daca dataExpirarii dupa prelungire a fost depasita contractul a expirat
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaContracteExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaContractePrelungite"]
        for i in range(len(randuriDeSters) - 1, -1, -1):
            ws.delete_rows(randuriDeSters[i])
        # 4.Pentru produsele in termen
        df_listaContracte = pd.read_excel(listaDB,
                                          sheet_name="ListaProduseInTermen")  # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        randuriDeSters = []
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            dataAsteptare = (dataExpirare + timedelta(days=5))
            if dataExpirare < now and dataAsteptare >= now:  # Produsul este in asteptare si trebuie mutat in lista cu ProduseInAsteptare si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaProduseInAsteptare"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
            elif dataAsteptare < now:  # Contractul a depasit si data asteptarii => este expirat
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaProduseExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaProduseInTermen"]
        for i in range(len(randuriDeSters) - 1, -1, -1):
            ws.delete_rows(randuriDeSters[i])
        # 2. Pentru produsele in asteptare
        randuriDeSters = []
        df_listaContracte = pd.read_excel(listaDB,
                                          sheet_name="ListaProduseInAsteptare")  # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            dataAsteptare = (dataExpirare + timedelta(days=5))
            if dataAsteptare < now:  # produsul este in asteptare si trebuie mutat in lista cu ProduseExpirate si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaProduseExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaProduseInAsteptare"]
        for i in range(len(randuriDeSters) - 1, -1, -1):
            ws.delete_rows(randuriDeSters[i])
        # 3. Pentru Produsele prelungite
        randuriDeSters = []
        df_listaContracte = pd.read_excel(listaDB,
                                              sheet_name="ListaProdusePrelungite")  # iau datele din tabel intrun panda.df pentru a accesa mai rapid data expirarii
        nrRanduri = len(df_listaContracte.index)
        for j in range(nrRanduri):
            dataExpirare = df_listaContracte.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            # dataAsteptare = (dataExpirare + timedelta(days=5))
            # if dataExpirare < now and dataAsteptare >= now:  # Contractul este in asteptare si trebuie mutat in lista cu ContracteInAsteptare si adaugat la lista cu indexul randului ce va trebui sters din lista contractelor in termen
            #     contract = df_listaContracte.loc[[j]]
            #     row = contract.values.tolist()
            #     ws = lista["ListaProduseInAsteptare"]
            #     ws.append(row[0])
            #     randuriDeSters.append(j + 2)
            if dataExpirare < now:  # Produsul a fost prelungit => data expirarii a fost extinsa , iar daca dataExpirarii dupa prelungire a fost depasita produsul a expirat
                contract = df_listaContracte.loc[[j]]
                row = contract.values.tolist()
                ws = lista["ListaProduseExpirate"]
                ws.append(row[0])
                randuriDeSters.append(j + 2)
        ws = lista["ListaProdusePrelungite"]
        for i in range(len(randuriDeSters) - 1, -1, -1):
            ws.delete_rows(randuriDeSters[i])
        lista.save(listaDB)

    def goToTabelContracteInTermen(self):
        self.ui.pushButton_lichideazaContractCuTermenDepasit.hide()
        self.ui.pushButton_incheieContract_6.show()
        self.ui.pushButton_incheieContract_7.show()

        self.ui.stackedWidget_4.setCurrentIndex(0)
        self.ui.tableWidget_contracteIncheiate.clearSelection()
        self.ui.tableWidget_contracteInTermen.clearSelection()
        self.ui.tableWidget_contracteInAsteptare.clearSelection()
        self.ui.tableWidget_contracteExpirate.clearSelection()
        self.ui.tableWidget_contractePrelungite.clearSelection()
    def goToTabelContracteInAsteptare(self):
        self.ui.pushButton_lichideazaContractCuTermenDepasit.hide()
        self.ui.pushButton_incheieContract_6.show()
        self.ui.pushButton_incheieContract_7.show()
        self.ui.stackedWidget_4.setCurrentIndex(1)
        self.ui.tableWidget_contracteIncheiate.clearSelection()
        self.ui.tableWidget_contracteInTermen.clearSelection()
        self.ui.tableWidget_contracteInAsteptare.clearSelection()
        self.ui.tableWidget_contracteExpirate.clearSelection()
        self.ui.tableWidget_contractePrelungite.clearSelection()
    def goToTabelContracteExpirate(self):
        self.ui.pushButton_finalizareIncheiereContract.hide()
        self.ui.pushButton_incheieContract_6.hide()
        self.ui.pushButton_incheieContract_7.hide()
        self.ui.pushButton_incheieContract_8.hide()
        self.ui.pushButton_lichideazaContractCuTermenDepasit.show()
        self.ui.stackedWidget_4.setCurrentIndex(2)
        self.ui.tableWidget_contracteIncheiate.clearSelection()
        self.ui.tableWidget_contracteInTermen.clearSelection()
        self.ui.tableWidget_contracteInAsteptare.clearSelection()
        self.ui.tableWidget_contracteExpirate.clearSelection()
        self.ui.tableWidget_contractePrelungite.clearSelection()
    def goToTabelContracteIncheiate(self):
        self.ui.pushButton_lichideazaContractCuTermenDepasit.hide()
        self.ui.pushButton_incheieContract_6.hide()
        self.ui.pushButton_incheieContract_7.hide()
        self.ui.stackedWidget_4.setCurrentIndex(3)
        self.ui.tableWidget_contracteIncheiate.clearSelection()
        self.ui.tableWidget_contracteInTermen.clearSelection()
        self.ui.tableWidget_contracteInAsteptare.clearSelection()
        self.ui.tableWidget_contracteExpirate.clearSelection()
        self.ui.tableWidget_contractePrelungite.clearSelection()
    def goToTabelContractePrelungite(self):
        self.ui.pushButton_incheieContract_6.show()
        self.ui.pushButton_incheieContract_7.show()
        self.ui.pushButton_lichideazaContractCuTermenDepasit.hide()
        self.ui.stackedWidget_4.setCurrentIndex(4)
        self.ui.tableWidget_contracteIncheiate.clearSelection()
        self.ui.tableWidget_contracteInTermen.clearSelection()
        self.ui.tableWidget_contracteInAsteptare.clearSelection()
        self.ui.tableWidget_contracteExpirate.clearSelection()
        self.ui.tableWidget_contractePrelungite.clearSelection()
    def goToGenereazaContract(self,listaDB):
        self.ui.radioButton_3.setChecked(False)
        self.ui.radioButton_4.setChecked(True)
        if self.ui.stackedWidget.currentIndex() == 3 :
            if not self.ui.tableWidget_4.selectedItems():
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Nu a fost selectat nici un client")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Ok:
                    print('OK clicked')
            else:
                CnpClient = self.ui.tableWidget_4.item(self.ui.tableWidget_4.currentRow(),2).text()
                self.cautaClientInDBDupaCNP(CnpClient,listaDB)
                self.ui.stackedWidget.setCurrentIndex(1)
        else:
            self.ui.stackedWidget.setCurrentIndex(1)

    def goToCautaClient(self):
        self.ui.stackedWidget.setCurrentIndex(3)
        self.ui.radioButton_2.setChecked(True)
        self.ui.radioButton.setChecked(False)
        self.ui.stackedWidget_2.setCurrentIndex(0)
        self.ui.stackedWidget_3.setCurrentIndex(3)
        self.ui.lineEdit_CautaNume.setText("")
        self.ui.lineEdit_CautaPrenume.setText("")
        self.ui.lineEdit_CautaCnp.setText("")
    def goHomePage(self,listaDB):
        # # print(self.ui.stackedWidget_4.currentIndex())
        if self.ui.stackedWidget_4.currentIndex() == 5:
            self.ui.stackedWidget_4.setCurrentIndex(0)
            self.ui.pushButton_incheieContract_7.show()
            self.ui.pushButton_incheieContract_6.show()
            self.ui.pushButton_incheieContract_8.hide()
            self.ui.pushButton_finalizareIncheiereContract.hide()
        elif self.ui.stackedWidget.currentIndex() == 6:
            self.ui.stackedWidget.setCurrentIndex(0)
        elif self.ui.stackedWidget.currentIndex() ==4:
            self.ui.stackedWidget.setCurrentIndex(0)
        else:
            self.nrContractelorAproapeDeTermen(listaDB)
            self.nrContracteInTermen(listaDB)
            self.produseInTermen(listaDB)
            self.ui.stackedWidget.setCurrentIndex(0)
            self.ui.lineEdit_NumeClient.setText("")
            self.ui.lineEdit_PrenumeClient.setText("")
            self.ui.lineEdit_SerieCI.setText("")
            self.ui.lineEdit_NumarCI.setText("")
            self.ui.lineEdit_CNP.setText("")
            self.ui.lineEdit_EliberatDe.setText("")
            self.ui.lineEdit_Adresa.setText((""))
            self.ui.lineEdit_Telefon.setText((""))
            self.ui.lineEdit_Email.setText((""))

    def goToContracteAproapeDeTermen(self):
        self.ui.stackedWidget.setCurrentIndex(4)
    def goToContracteInTermen(self):
        self.ui.pushButton_lichideazaContractCuTermenDepasit.hide()
        self.ui.pushButton_incheieContract_8.hide()
        self.ui.pushButton_finalizareIncheiereContract.hide()
        self.ui.stackedWidget.setCurrentIndex(5)
        self.ui.stackedWidget_4.setCurrentIndex(0)
    def goToProduseInTermen(self):
        self.ui.stackedWidget.setCurrentIndex(6)
    def goToClientNegasit(self):
        self.ui.stackedWidget_3.setCurrentIndex(1)
    def goToClientGasit(self):
        self.ui.stackedWidget_3.setCurrentIndex(0)

    def goToMaiMultiClienti(self):
        self.ui.stackedWidget_3.setCurrentIndex(2)
    def nrContractelorAproapeDeTermen(self,listaDB):
        print("START nrContractelorAproapeDeTermen")
        now = date.today()
        listaDB = listaDB
        dataPesteTreiZile = (date.today() + timedelta(days=3))
        dataPesteDouaZile = (date.today() + timedelta(days=2))
        dataPesteOZi = (date.today() + timedelta(days=1))
        listaContracte = ["ListaContracteInTermen","ListaContractePrelungite","ListaContracteInAsteptare"]
        df_ToateContractele = pd.DataFrame()
        for i in range(len(listaContracte)):
            df = pd.read_excel(listaDB,
                                              sheet_name=listaContracte[i])
            df_ToateContractele = df_ToateContractele.append(df,ignore_index=True)
        nrRanduri = len(df_ToateContractele.index)
        listaContracte = [0, "", "", "", ""]
        # listaContracte[0] = nr de contracte care expira in urmatoarele 3 zile
        # listaContracte[1] = contractele care expira azi
        # listaContracte[2] = contractele care expira maine
        # listaContracte[3] = contractele care expira peste 2 zile
        # listaContracte[4] =  contractele care expira peste 3 zile
        for j in range(nrRanduri):
            dataExpirare = df_ToateContractele.iloc[j]['Data expirare']
            dataExpirare = datetime.strptime(dataExpirare, '%d.%m.%Y').date()
            if dataExpirare == now:
                listaContracte[1] = listaContracte[1] +  str(df_ToateContractele.iloc[j]['Nr. Contract']) + ","
                listaContracte[0] +=1
            elif dataExpirare == dataPesteOZi:
                listaContracte[2] = listaContracte[2] +  str(df_ToateContractele.iloc[j]['Nr. Contract']) + ","
                listaContracte[0] += 1
            elif dataExpirare == dataPesteDouaZile:
                listaContracte[3] = listaContracte[3] +  str(df_ToateContractele.iloc[j]['Nr. Contract']) + ","
                listaContracte[0] += 1
            elif dataExpirare == dataPesteTreiZile:
                listaContracte[4] = listaContracte[4] +  str(df_ToateContractele.iloc[j]['Nr. Contract']) + ","
                listaContracte[0] += 1
        listaContracte[1] = listaContracte[1][:-1]
        listaContracte[2] = listaContracte[2][:-1]
        listaContracte[3] = listaContracte[3][:-1]
        listaContracte[4] = listaContracte[4][:-1]
        self.ui.label_numarCercStanga.setText(str(listaContracte[0]))
        print("STOP nrContractelorAproapeDeTermen - return:" ,listaContracte)
        return listaContracte
    def adaugaValoriInTabelContracte(self,listaDB,listaContracte,tabel):
        print("START adaugaValoriInTabelContracte")
        df = pd.read_excel(listaDB, sheet_name=listaContracte)
        tabel.setRowCount(len(df.index))
        if tabel == self.ui.tableWidget_contracteInTermen or tabel == self.ui.tableWidget_contracteInAsteptare:
            tabel.setColumnWidth(0, 79)
            tabel.setColumnWidth(1, 300)
            tabel.setColumnWidth(2, 300)
            tabel.setColumnWidth(3, 131)
            tabel.setColumnWidth(4, 211)
            tabel.setColumnWidth(5, 86)
            tabel.setColumnWidth(6, 93)
            tabel.setColumnWidth(7, 125)
            tabel.setColumnWidth(8, 66)
            tabel.setColumnWidth(9, 113)
        row = 0
        # Verific daca se populeaza lista contractelor in termen, deoarece aceasta lista nu contine contracte prelungite
        # La fel si lista contractelor in asteptare, unde pot intra doar contractele care nu au mai fost prelungite
        if listaContracte == "ListaContracteInTermen" or listaContracte == "ListaContracteInAsteptare":
            print(listaContracte)
            for i in range(len(df.index)):
                tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Nume"]))
                tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Prenume"])))
                tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Telefon"])))
                tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Email"])))
                tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma imprumutata"])))
                tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision"])))
                tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit"])))
                row += 1
        # elif daca se populeaza lista contractelor prelungite, acestea prin definitie sunt prelungite si urmeaza tiparul de mai jos
        elif listaContracte == "ListaContractePrelungite":
            tabel.setColumnWidth(0, 79)
            tabel.setColumnWidth(1, 250)
            tabel.setColumnWidth(2, 250)
            tabel.setColumnWidth(3, 131)
            tabel.setColumnWidth(4, 211)
            tabel.setColumnWidth(5, 86)
            tabel.setColumnWidth(6, 93)
            tabel.setColumnWidth(7, 125)
            tabel.setColumnWidth(8, 66)
            tabel.setColumnWidth(9, 113)
            tabel.setColumnWidth(10, 132)
            tabel.setColumnWidth(11, 102)
            tabel.setColumnWidth(12, 112)
            tabel.setColumnWidth(13, 130)
            tabel.setColumnWidth(14, 207)
            for i in range(len(df.index)):
                tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Nume"]))
                tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Prenume"])))
                tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Telefon"])))
                tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Email"])))
                tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma imprumutata"])))
                tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision"])))
                tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit"])))
                tabel.setItem(row, 10, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare initiala"])))
                tabel.setItem(row, 11, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data prelungire"])))
                tabel.setItem(row, 12, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Zile prelungite"])))
                tabel.setItem(row, 13, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision prelungire"])))
                tabel.setItem(row, 14, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit dupa prelungire"])))
                row += 1
        # Daca nu se populeaza lista contractelor in termen sau cea a contractelor prelungite sau cea a contractelor in asteptare, verific daca coloana "Data expirare initiala" nu contine valori => contractul nu este unul prelungit
        else:
            tabel.setColumnWidth(0, 79)
            tabel.setColumnWidth(1, 250)
            tabel.setColumnWidth(2, 250)
            tabel.setColumnWidth(3, 131)
            tabel.setColumnWidth(4, 211)
            tabel.setColumnWidth(5, 86)
            tabel.setColumnWidth(6, 93)
            tabel.setColumnWidth(7, 125)
            tabel.setColumnWidth(8, 66)
            tabel.setColumnWidth(9, 113)
            tabel.setColumnWidth(10, 132)
            tabel.setColumnWidth(11, 102)
            tabel.setColumnWidth(12, 112)
            tabel.setColumnWidth(13, 130)
            tabel.setColumnWidth(14, 207)
            for i in range(len(df.index)):
                if pd.isnull(df.iloc[i]["Data expirare initiala"]):
                    tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                    tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Nume"]))
                    tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Prenume"])))
                    tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Telefon"])))
                    tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Email"])))
                    tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                    tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                    tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma imprumutata"])))
                    tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision"])))
                    tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit"])))
                    tabel.setItem(row, 10, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 11, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 12, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 13, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 14, QtWidgets.QTableWidgetItem("-"))
                    row += 1
                else: # Contractul este unul care a fost prelungit
                    tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                    tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Nume"]))
                    tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Prenume"])))
                    tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Telefon"])))
                    tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Email"])))
                    tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                    tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                    tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma imprumutata"])))
                    tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision"])))
                    tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit"])))
                    tabel.setItem(row, 10, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare initiala"])))
                    tabel.setItem(row, 11, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data prelungire"])))
                    tabel.setItem(row, 12, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Zile prelungite"])))
                    tabel.setItem(row, 13, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision prelungire"])))
                    tabel.setItem(row, 14, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit dupa prelungire"])))
                    row += 1
        print("STOP adaugaValoriInTabelContracte- no return:")

    def PopulareTabelContracte(self,listaDB):
        print("START PopulareTabelContracte")
        listaContracte = ["ListaContracteInTermen","ListaContractePrelungite","ListaContracteExpirate","ListaContracteLichidate","ListaContracteInAsteptare"]
        listaTabele = [self.ui.tableWidget_contracteInTermen,self.ui.tableWidget_contractePrelungite,self.ui.tableWidget_contracteExpirate,self.ui.tableWidget_contracteIncheiate,self.ui.tableWidget_contracteInAsteptare]
        for i in range(len(listaContracte)):
            self.adaugaValoriInTabelContracte(listaDB, listaContracte[i], listaTabele[i])
        print("STOP PopulareTabelContracte- no return:")
    def adaugaValoriInTabelProduse(self,listaDB,listaContracte,tabel):
        print("START adaugaValoriInTabelProduse")
        df = pd.read_excel(listaDB, sheet_name=listaContracte)
        tabel.setRowCount(len(df.index))
        row = 0
        # Verific daca se populeaza lista produselor in termen, deoarece aceasta lista nu contine produse prelungite
        # La fel si lista produselor in asteptare, unde pot intra doar contractele care nu au mai fost prelungite
        if listaContracte == "ListaProduseInTermen" or listaContracte == "ListaProduseInAsteptare":
            # Setare dimensiune coloane tabel
            tabel.setColumnWidth(0, 353)
            tabel.setColumnWidth(1, 222)
            tabel.setColumnWidth(2, 80)
            tabel.setColumnWidth(3, 80)
            tabel.setColumnWidth(4, 90)
            tabel.setColumnWidth(5, 125)
            tabel.setColumnWidth(6, 134)
            self.ui.tableWidget_produseInTermen.setColumnWidth(7, 106)
            for i in range(len(df.index)):
                tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Observatii"]))
                tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                row += 1
        # elif daca se populeaza lista contractelor prelungite, acestea prin definitie sunt prelungite si urmeaza tiparul de mai jos
        elif listaContracte == "ListaProdusePrelungite":
            tabel.setColumnWidth(0, 353)
            tabel.setColumnWidth(1, 222)
            tabel.setColumnWidth(2, 80)
            tabel.setColumnWidth(3, 80)
            tabel.setColumnWidth(4, 90)
            tabel.setColumnWidth(5, 125)
            tabel.setColumnWidth(6, 65)
            tabel.setColumnWidth(7, 106)
            tabel.setColumnWidth(8, 132)
            tabel.setColumnWidth(9, 102)
            tabel.setColumnWidth(10, 114)
            tabel.setColumnWidth(11, 131)
            tabel.setColumnWidth(12, 207)
            for i in range(len(df.index)):
                tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Observatii"]))
                tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare initiala"])))
                tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data prelungire"])))
                tabel.setItem(row, 10, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Zile prelungite"])))
                tabel.setItem(row, 11, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision prelungire"])))
                tabel.setItem(row, 12, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit dupa prelungire"])))
                row += 1
        elif listaContracte == "ListaProduseLichidateClient":
            tabel.setColumnWidth(0, 353)
            tabel.setColumnWidth(1, 222)
            tabel.setColumnWidth(2, 80)
            tabel.setColumnWidth(3, 80)
            tabel.setColumnWidth(4, 90)
            tabel.setColumnWidth(5, 125)
            tabel.setColumnWidth(6, 65)
            tabel.setColumnWidth(7, 106)
            tabel.setColumnWidth(8, 132)
            tabel.setColumnWidth(9, 102)
            tabel.setColumnWidth(10, 114)
            tabel.setColumnWidth(11, 131)
            tabel.setColumnWidth(12, 207)
            tabel.setColumnWidth(13, 207)
            for i in range(len(df.index)):
                tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Observatii"]))
                tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare initiala"])))
                tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data prelungire"])))
                tabel.setItem(row, 10, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Zile prelungite"])))
                tabel.setItem(row, 11, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision prelungire"])))
                tabel.setItem(row, 12,
                              QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit dupa prelungire"])))
                tabel.setItem(row, 13,
                              QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data lichidare"])))
                row += 1
        # Daca nu se populeaza lista produselor in termen sau cea a produselor prelungite, verific daca coloana "Data expirare initiala" nu contine valori => produsul nu este unul prelungit
        else:
            tabel.setColumnWidth(0, 353)
            tabel.setColumnWidth(1, 222)
            tabel.setColumnWidth(2, 80)
            tabel.setColumnWidth(3, 80)
            tabel.setColumnWidth(4, 90)
            tabel.setColumnWidth(5, 125)
            tabel.setColumnWidth(6, 65)
            tabel.setColumnWidth(7, 106)
            tabel.setColumnWidth(8, 132)
            tabel.setColumnWidth(9, 102)
            tabel.setColumnWidth(10, 114)
            tabel.setColumnWidth(11, 131)
            tabel.setColumnWidth(12, 207)
            for i in range(len(df.index)):
                if pd.isnull(df.iloc[i]["Data expirare initiala"]):
                    tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                    tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Observatii"]))
                    tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                    tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                    tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                    tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                    tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                    tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                    tabel.setItem(row, 8, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 9, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 10, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 11, QtWidgets.QTableWidgetItem("-"))
                    tabel.setItem(row, 12,
                                  QtWidgets.QTableWidgetItem("-"))
                    row += 1
                else:  # Contractul este unul care a fost prelungit
                    tabel.setItem(row, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                    tabel.setItem(row, 1, QtWidgets.QTableWidgetItem(df.iloc[i]["Observatii"]))
                    tabel.setItem(row, 2, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                    tabel.setItem(row, 3, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                    tabel.setItem(row, 4, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                    tabel.setItem(row, 5, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                    tabel.setItem(row, 6, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                    tabel.setItem(row, 7, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                    tabel.setItem(row, 8, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare initiala"])))
                    tabel.setItem(row, 9, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data prelungire"])))
                    tabel.setItem(row, 10, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Zile prelungite"])))
                    tabel.setItem(row, 11, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision prelungire"])))
                    tabel.setItem(row, 12,
                                  QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit dupa prelungire"])))
                    row += 1
        print("STOP adaugaValoriInTabelProduse- no return:")
    def PopulareTabelProduse(self,listaDB):
        print("START PopulareTabelProduse")
        listaProduse = ["ListaProduseInTermen", "ListaProduseInAsteptare", "ListaProdusePrelungite",
                          "ListaProduseExpirate", "ListaProduseLichidateClient"]
        listaTabele = [self.ui.tableWidget_produseInTermen, self.ui.tableWidget_produseInAsteptare,
                       self.ui.tableWidget_produsePrelungite, self.ui.tableWidget_produseExpirate,
                       self.ui.tableWidget_produseLichidateClient]
        for i in range(len(listaProduse)):
            print(listaProduse[i])
            self.adaugaValoriInTabelProduse(listaDB, listaProduse[i], listaTabele[i])
        print("STOP PopulareTabelProduse- no return:")
    def PopulareTabelGestiuneVanzariProduseAmanet(self,listaDB):
        print("START PopulareTabelGestiuneVanzariProduseAmanet")
        tabel = self.ui.tableWidget_produseLichidateClient_5
        tabel.setColumnWidth(0, 385 )
        tabel.setColumnWidth(1, 98)
        df = pd.read_excel(listaDB,sheet_name="ListaProduseLichidateAmanet")
        tabel.setRowCount(len(df.index))
        for i in range(len(df.index)):
            tabel.setItem(i, 0, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
            tabel.setItem(i, 1, QtWidgets.QTableWidgetItem(str(df.iloc[i]["Pret vanzare"])))
        print("STOP PopulareTabelGestiuneVanzariProduseAmanet- no return:")
    def PopulareTabelGestiuneVanzareProduseFurnizor(self,listaDB):
        print("START PopulareTabelGestiuneVanzareProduseFurnizor")
        df = pd.read_excel(listaDB, sheet_name="ListaProduseFurnizori")
        tabel = self.ui.tableWidget_produseLichidateClient_3
        tabel.setRowCount(len(df.index))
        for i in range(len(df.index)):
            denumireProdus = df.iloc[i]["Denumire Produs"]
            codProdus = df.iloc[i]["Cod produs"]
            pret = df.iloc[i]["Valoare la vanzare"]
            pret = str(pret)
            pret = pret.split(",")
            pret = int(pret[0])
            if not pd.isnull(df.iloc[i]["Cantitate vanduta"]):
                cantitateVanduta = df.iloc[i]["Cantitate vanduta"]
                cantitateVanduta = str(cantitateVanduta)
                cantitateVanduta = cantitateVanduta.split(",")
                for j in range(len(cantitateVanduta)):
                    cantitateVanduta[j] = int(cantitateVanduta[j])
                cantitateVanduta = sum(cantitateVanduta)
            else:
                cantitateVanduta = 0
            cantitateInitiala = df.iloc[i]["Cantitate intrare"]
            cantitateInitiala = str(cantitateInitiala)
            cantitateInitiala = cantitateInitiala.split(",")
            for j in range(len(cantitateInitiala)):
                cantitateInitiala[j] = int(cantitateInitiala[j])
            cantitateInitiala = sum(cantitateInitiala)
            stoc = cantitateInitiala - cantitateVanduta
            tabel.setItem(i,0,QtWidgets.QTableWidgetItem(denumireProdus))
            tabel.setItem(i,1,QtWidgets.QTableWidgetItem(codProdus))
            tabel.setItem(i,2,QtWidgets.QTableWidgetItem(str(pret)))
            tabel.setItem(i,3,QtWidgets.QTableWidgetItem(str(stoc)))
        print("STOP PopulareTabelGestiuneVanzareProduseFurnizor- no return:")
    def verificaDateProdusFurnizor(self):
        tabel = self.ui.tableWidget_produseLichidateClient_8
        nrProduse = self.ui.tableWidget_produseLichidateClient_8.rowCount()
        ok = 0
        if nrProduse == 0:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu au fost introduse produse pentru a fi adaugate in sistem")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            return False
        else:
            print("nrProduse:",nrProduse)
            for i in range(nrProduse):
                # numeProdus = tabel.item(i,0).text()
                # codProdus = tabel.item(i,1).text()
                # cantitateFactura = tabel.item(i,2).text()
                # cantitatePrimita = tabel.item(i,3).text()
                # pretAchizitie = tabel.item(i,4).text()
                # pretVanzare = tabel.item(i,5).text()
                print("nrRand :",i)
                if tabel.item(i,0) is None or tabel.item(i,0).text()=="" :
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdus numele produsului de pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif tabel.item(i,1) is None:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdus codul produsului de pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif tabel.item(i,2) is None:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdusa cantitate din factura pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif  not tabel.item(i,2).text().isdigit():
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Cantitate din factura de pe randul " + str(i + 1) + " trebuie sa contina doar cifre\nex:100")
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif tabel.item(i,3) is None:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdusa cantitate primita pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif not tabel.item(i,3).text().isdigit():
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText(
                        "Cantitate primita de pe randul " + str(i + 1) + " trebuie sa contina doar cifre\nex:100")
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif tabel.item(i,4) is None:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdus pretul de achizitie pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif not tabel.item(i,4).text().isdigit():
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Pretul de achizitie de pe randul " + str(i + 1) + " trebuie sa contina doar cifre\nex:100")
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif tabel.item(i,5) is None:
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Nu a fost introdus pretul de vanzare pe randul "+ str(i+1))
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif not tabel.item(i,5).text().isdigit():
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText(
                        "Pretul de vanzare de pe randul " + str(i + 1) + " trebuie sa contina doar cifre\nex:100")
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                        print('OK clicked')
                        tabel.selectRow(i)
                    return False
                elif i == int(nrProduse)-1:
                    return True
    def adaugaProdusFurnizor(self):
        nrRanduri = self.ui.tableWidget_produseLichidateClient_8.rowCount()
        self.ui.tableWidget_produseLichidateClient_8.setRowCount(nrRanduri+1)
    def finalizareAdaugareProduseFurnizor(self,listaDB):
        if self.verificaDateProdusFurnizor():
            nrFactura = self.ui.lineEdit_PrenumeClient_7.text()
            numeFurnizor = self.ui.lineEdit_SerieCI_5.text()
            if nrFactura=="" or not nrFactura.isdigit():
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Nu a fost introdus numarul facturii")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Ok:
                    print('OK clicked')
            elif numeFurnizor =="":
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Nu a fost introdus numele furnizorului")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Ok:
                    print('OK clicked')
            else:
                df = pd.read_excel(listaDB, sheet_name="ListaProduseFurnizori")
                nrProduse = self.ui.tableWidget_produseLichidateClient_8.rowCount()
                # Adaug produsele in DB
                lista = load_workbook(listaDB)
                ws = lista["ListaProduseFurnizori"]
                for i in range(nrProduse):
                    codProdus = self.ui.tableWidget_produseLichidateClient_8.item(i,1).text()
                    pretVanzare = self.ui.tableWidget_produseLichidateClient_8.item(i,5).text()
                    cantitateIntrare = self.ui.tableWidget_produseLichidateClient_8.item(i,3).text()
                    numeProdus = self.ui.tableWidget_produseLichidateClient_8.item(i,0).text()
                    for j in range(len(df.index)):
                        codProdusDB = df.iloc[j]["Cod produs"]
                        if codProdus == codProdusDB:
                            ws['C'+ str(j+2)] = pretVanzare + ',' + str(ws['C'+ str(j+2)].value)
                            ws['D' + str(j+2)] = date.today().strftime("%d.%m.%Y") + ","+ str(ws['D' + str(j+2)].value)
                            ws['E'+str(j+2)] = cantitateIntrare + ","+ str(ws['E'+str(j+2)].value)
                        else:
                            row =[]
                            row.append(numeProdus)
                            row.append(codProdus)
                            row.append(pretVanzare)
                            row.append(date.today().strftime("%d.%m.%Y"))
                            row.append(cantitateIntrare)
                            ws.append(row)
                lista.save(listaDB)
                self.PopulareTabelGestiuneVanzareProduseFurnizor(listaDB)
                # Generare NIR pentru produsele adaugate
                nrNir = self.genereazaNrNir(lista,listaDB)
                ws = lista["TemplateNIR"]
                ws['I3'] = "NI-" + str(nrNir)
                ws['I4'] = date.today().strftime('%d.%m.%Y')
                ws['B6'] = "Subsemnatii, membrii ai comisiei de receptie, am receptionat valorile materiale furnizate de: "+ numeFurnizor +" pe baza documentelor insotitoare: Factura NR: "+ nrFactura + " ,constatand:"
                ws.merge_cells(start_row=6, start_column=2, end_row=7, end_column=14)
                ws.insert_rows(11, amount=nrProduse)
                totalAchizitie = 0
                totalAdaos = 0
                totalVanzare = 0
                totalProduse = 0
                for i in range(nrProduse):
                    ws.merge_cells(start_row=11 + i, start_column=2, end_row=11 + i, end_column=4)
                    ws['A' + str(11 + i)] = i + 1
                    pretAchizitie = int(self.ui.tableWidget_produseLichidateClient_8.item(i, 4).text())
                    pretVanzare = int(self.ui.tableWidget_produseLichidateClient_8.item(i, 5).text())
                    cantitateIntrare = self.ui.tableWidget_produseLichidateClient_8.item(i, 3).text()
                    numeProdus = self.ui.tableWidget_produseLichidateClient_8.item(i, 0).text()
                    codProdus = self.ui.tableWidget_produseLichidateClient_8.item(i, 1).text()
                    numeProdus = numeProdus + " - "+codProdus
                    cantitateDoc = self.ui.tableWidget_produseLichidateClient_8.item(i, 2).text()
                    ws['B' + str(11 + i)] = numeProdus
                    ws['E' + str(11 + i)] = "B.C"
                    ws['F' + str(11 + i)] = cantitateDoc
                    ws['G' + str(11 + i)] = cantitateIntrare
                    ws['H' + str(11 + i)] = pretAchizitie
                    ws['I' + str(11 + i)] = int(ws['G' + str(11 + i)].value) * int(ws['H' + str(11 + i)].value)
                    ws['M' + str(11 + i)] = pretVanzare
                    ws['N' + str(11 + i)] = pretVanzare * int(ws['G' + str(11 + i)].value)
                    adaos = int(pretVanzare) - int(pretAchizitie)
                    adaosProcentual = ((int(pretVanzare) - int(pretAchizitie)) / int(pretAchizitie)) * 100
                    adaosProcentual = round(adaosProcentual, 2)
                    ws['J' + str(11 + i)] = adaosProcentual
                    ws['K' + str(11 + i)] = adaos
                    ws['L' + str(11 + i)] = adaos*int(ws['G' + str(11 + i)].value)
                    totalAchizitie = totalAchizitie + ws['I' + str(11 + i)].value
                    totalAdaos = totalAdaos + ws['L' + str(11 + i)].value
                    totalVanzare = totalVanzare + ws['N' + str(11 + i)].value
                    totalProduse = totalProduse + int(cantitateIntrare)
                    self.styleTabel(ws, 11 + i - 1, 'A')
                    self.styleTabel(ws, 11 + i - 1, 'B')
                    self.styleTabel(ws, 11 + i - 1, 'C')
                    self.styleTabel(ws, 11 + i - 1, 'D')
                    self.styleTabel(ws, 11 + i - 1, 'E')
                    self.styleTabel(ws, 11 + i - 1, 'F')
                    self.styleTabel(ws, 11 + i - 1, 'G')
                    self.styleTabel(ws, 11 + i - 1, 'H')
                    self.styleTabel(ws, 11 + i - 1, 'I')
                    self.styleTabel(ws, 11 + i - 1, 'M')
                    self.styleTabel(ws, 11 + i - 1, 'N')
                    self.styleTabel(ws, 11 + i - 1, 'J')
                    self.styleTabel(ws, 11 + i - 1, 'K')
                    self.styleTabel(ws, 11 + i - 1, 'L')
                ws['G' + str(11 + nrProduse)] = totalProduse
                ws['H' + str(11 + nrProduse)] = totalAchizitie
                ws['J' + str(11 + nrProduse)] = totalAdaos
                ws['M' + str(11 + nrProduse)] = totalVanzare
                ws.merge_cells(start_row=int(11 + nrProduse), start_column=8, end_row=int(11 + nrProduse), end_column=9)
                ws.merge_cells(start_row=int(11 + nrProduse), start_column=10, end_row=int(11 + nrProduse), end_column=12)
                ws.merge_cells(start_row=int(11 + nrProduse), start_column=13, end_row=int(11 + nrProduse), end_column=14)
                for i in range(nrProduse - 1, -1, -1):
                    self.ui.tableWidget_produseLichidateClient_8.removeRow(i)
                self.ui.lineEdit_PrenumeClient_7.setText("")
                self.ui.lineEdit_SerieCI_5.setText("")
                currentDir = os.getcwd()
                numeNir = currentDir + "\\NIR-uri\\" + "A-" + str(nrNir) + ".xlsx"
                lista.save(numeNir)
                os.startfile(numeNir)
                lista.save(listaDB)
    def vindeProdusFurnizor(self,listaDB):
        if not self.ui.tableWidget_produseLichidateClient_3.selectedItems():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu a fost selectat nici un produs")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            bucVandute, okPressed = QInputDialog.getInt(self, "Numar buc vandute","Introdu numarul de bucati vandute: ", 1, 0, 100, 1)
            if okPressed:
                df_produse = pd.read_excel(listaDB, sheet_name="ListaProduseFurnizori")
                lista = load_workbook(listaDB)
                for i in range(len(df_produse.index)):
                    randSelectat = self.ui.tableWidget_produseLichidateClient_3.currentRow()
                    if self.ui.tableWidget_produseLichidateClient_3.item(randSelectat,1).text() == df_produse.iloc[i]["Cod produs"]:
                        ws = lista["ListaProduseFurnizori"]
                        ws['F' + str(i + 2)] = date.today().strftime("%d.%m.%Y") + ',' + str(ws['F' + str(i + 2)].value)
                        ws['G' + str(i + 2)] =str(bucVandute) + "," + str(ws['G' + str(i + 2)].value)
                        lista.save(listaDB)
                        self.PopulareTabelGestiuneVanzareProduseFurnizor(listaDB)
    def selecteazaContractInTabel(self):
        print("START selecteazaContractInTabel")
        gasit = 0
        valoareCautata = self.ui.lineEdit_numarContractDePrelungit_4.text()
        listaTabele = [self.ui.tableWidget_contracteInTermen,self.ui.tableWidget_contracteInAsteptare,self.ui.tableWidget_contracteExpirate,self.ui.tableWidget_contracteIncheiate,self.ui.tableWidget_contractePrelungite]
        for i in range(len(listaTabele)):
            nrRows = listaTabele[i].rowCount()
            listaTabele[i].clearSelection()
            for j in range(nrRows):
                if listaTabele[i].item(j, 0):
                    if valoareCautata == listaTabele[i].item(j, 0).text():
                        if i == 0:
                            self.goToTabelContracteInTermen()
                        elif i == 1:
                            self.goToTabelContracteInAsteptare()
                        elif i ==2:
                            self.goToTabelContracteExpirate()
                        elif i ==3:
                            self.goToTabelContracteIncheiate()
                        elif i ==4:
                            self.goToTabelContractePrelungite()
                        listaTabele[i].selectRow(j)
                        gasit = 1
        if gasit == 0:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Contractul nu exista")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        print("STOP selecteazaContractInTabel- no return:")

    def selecteazaContractInTabelProduse(self):
        print("START selecteazaContractInTabel")
        gasit = 0
        valoareCautata = self.ui.lineEdit_numarContractDePrelungit_5.text()
        listaTabele = [self.ui.tableWidget_produseInTermen, self.ui.tableWidget_produseInAsteptare,
                       self.ui.tableWidget_produsePrelungite, self.ui.tableWidget_produseExpirate,
                       self.ui.tableWidget_produseLichidateClient]
        for i in range(len(listaTabele)):
            nrRows = listaTabele[i].rowCount()
            listaTabele[i].clearSelection()
            for j in range(nrRows):
                if listaTabele[i].item(j, 2):
                    if valoareCautata == listaTabele[i].item(j, 2).text():
                        if i == 0:
                            self.goToProduseInTermen()
                        elif i == 1:
                            self.ui.stackedWidget_5.setCurrentIndex(1)
                        elif i == 2:
                            self.ui.stackedWidget_5.setCurrentIndex(2)
                        elif i == 3:
                            self.ui.stackedWidget_5.setCurrentIndex(3)
                        elif i == 4:
                            self.ui.stackedWidget_5.setCurrentIndex(4)
                        listaTabele[i].selectRow(j)
                        gasit = 1
        if gasit == 0:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Contractul nu exista")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        print("STOP selecteazaContractInTabel- no return:")
    def lichideazaContractCuTermenDepasit(self,listaDB):
        if not self.ui.tableWidget_contracteExpirate.selectedItems():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu a fost selectat nici un contract")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            indexRand = self.ui.tableWidget_contracteExpirate.selectedIndexes()
            nrContract = self.ui.tableWidget_contracteExpirate.item(self.ui.tableWidget_contracteExpirate.currentRow(), 0).text()
            # Preia datele contractului din DB
            df_contract = pd.read_excel(listaDB, sheet_name="ListaContracteExpirate")
            randContract = 0
            for i in range(len(df_contract.index)):
                if int(nrContract) == df_contract.iloc[i]["Nr. Contract"]:
                    df_contract = df_contract.loc[i]
                    randContract = i + 2
                    break
            # Preia date produse din contractul selectat
            df_toateProdusele = pd.read_excel(listaDB, sheet_name="ListaProduseExpirate")  # Iau toate produsele
            df_produse = pd.DataFrame(
                columns=["Denumire Produs", "Observatii", "Valoare Evaluata", "Suma Imprumutata", "Comision(%/zi)",
                         "Comision (lei)", "Data intrare in amanet", "Data expirare", "Total de restituit (lei)",
                         "Nr. Contract", "Data expirare initiala", "Data prelungire", "Nr. Zile prelungite",
                         "Comision prelungire", "Total de restituit dupa prelungire", "Data lichidare"])
            randuriProduseDeSters = []
            for i in range(len(df_toateProdusele.index)):
                if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                    df_produse = df_produse.append(df_toateProdusele.loc[i])
                    randuriProduseDeSters.append(i + 2)
            row = df_contract.values.tolist()
            row.append(date.today().strftime('%d.%m.%Y'))
            row.append("AMANET")
            lista = load_workbook(listaDB)
            ws = lista["ListaContracteLichidate"]
            ws.append(row)
            nrRanduri = self.ui.tableWidget_contracteIncheiate.rowCount()
            self.ui.tableWidget_contracteIncheiate.setRowCount(nrRanduri + 1)
            for i in range(len(row)):
                self.ui.tableWidget_contracteIncheiate.setItem(nrRanduri, i,
                                                               QtWidgets.QTableWidgetItem(str(row[i])))
            self.ui.tableWidget_contracteExpirate.removeRow(self.ui.tableWidget_contracteExpirate.currentRow())
            # sterg contractul din lista cu contracte in termen
            ws = lista["ListaContracteExpirate"]
            ws.delete_rows(randContract)
            row = df_produse.values.tolist()
            for i in range(len(row) - 1, -1, -1):
                row[i][15] = date.today().strftime('%d.%m.%Y')
                ws = lista["ListaProduseLichidateAmanet"]
                ws.append(row[i])
                ws = lista["ListaProduseExpirate"]
                ws.delete_rows(randuriProduseDeSters[i])
                print(row)
                # adaug produsul in tabelul din gestiunea de vanzari pentru a modifica detaliile produsului
                nrRanduri = self.ui.tableWidget_produseLichidateClient_6.rowCount()
                self.ui.tableWidget_produseLichidateClient_6.setRowCount(nrRanduri+1)
                self.ui.tableWidget_produseLichidateClient_6.setItem(nrRanduri, 0,QtWidgets.QTableWidgetItem(str(nrContract)))
                self.ui.tableWidget_produseLichidateClient_6.setItem(nrRanduri, 1,QtWidgets.QTableWidgetItem(str(row[i][0])))
                self.ui.tableWidget_produseLichidateClient_6.setItem(nrRanduri, 2,QtWidgets.QTableWidgetItem(str(row[i][1])))
                self.ui.tableWidget_produseLichidateClient_6.setItem(nrRanduri, 3,QtWidgets.QTableWidgetItem(str(row[i][3])))
                self.ui.tableWidget_produseLichidateClient_6.setItem(nrRanduri, 4,QtWidgets.QTableWidgetItem(str(row[i][8])))
            lista.save(listaDB)
            self.PopulareTabelContracte(listaDB)
            self.PopulareTabelProduse(listaDB)
            self.ui.pushButton_goHome_25.hide()
            self.goToGestiuneVanzari()
    def adaugaPretDeVanzareProdusAmanet(self,listaDB):
        nrProduse = self.ui.tableWidget_produseLichidateClient_6.rowCount()
        lista = load_workbook(listaDB)
        ws = lista["ListaProduseLichidateAmanet"]
        df = pd.read_excel(listaDB,sheet_name="ListaProduseLichidateAmanet")
        for j in range(nrProduse):
            if self.ui.tableWidget_produseLichidateClient_6.item(j,
                                                                 5) is None or self.ui.tableWidget_produseLichidateClient_6.item(
                    j, 5).text() == "":
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Pretul de vanzare nu a fost specificat")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Ok:
                    print('OK clicked')
                    self.ui.tableWidget_produseLichidateClient_6.selectRow(j)
                return False
            elif not self.ui.tableWidget_produseLichidateClient_6.item(j, 5).text().isdigit():
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Pretul de vanzare trebuie sa contina doar cifre\nex:5000")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Ok:
                    print('OK clicked')
                    self.ui.tableWidget_produseLichidateClient_6.selectRow(j)
                return False
            else:
                pret =  self.ui.tableWidget_produseLichidateClient_6.item(j, 5).text()
                numeProdus = self.ui.tableWidget_produseLichidateClient_6.item(j, 1).text()
                for i in range(len(df.index)):
                    if numeProdus == df.iloc[i]["Denumire Produs"]:
                        ws['R' + str(i + 2)] = pret
                lista.save(listaDB)
                return True
    def genereazaNrNir(self,lista,listaDB):
        print("START genereazaNrNir")
        ws = lista["ListaNIR"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        ws['B' + str(maxRow + 1)] = date.today().strftime("%d.%m.%Y")
        self.styleTabel(ws, maxRow, 'A')
        self.styleTabel(ws, maxRow, 'B')
        lista.save(listaDB)
        print("STOP genereazaNrNir- return:", maxRow)
        return maxRow
    def adaugaProdusAmanetLaVanzare(self,listaDB):
        nrProduse = self.ui.tableWidget_produseLichidateClient_6.rowCount()
        if self.adaugaPretDeVanzareProdusAmanet(listaDB):
            lista = load_workbook(listaDB)
            nrNir = self.genereazaNrNir(lista,listaDB)
            ws = lista["TemplateNIR"]
            ws['I3'] = "NI-" + str(nrNir)
            ws['I4'] = date.today().strftime('%d.%m.%Y')
            nrContract = self.ui.tableWidget_produseLichidateClient_6.item(0,0).text()
            ws['B6'] = "Subsemnatii, membrii ai comisiei de receptie, am receptionat valorile materiale furnizate pe baza documentelor insotitoare: Contract numarul "+ nrContract+ " constatand:"
            ws.merge_cells(start_row=6, start_column=2, end_row=7, end_column=14)
            ws.insert_rows(11,amount=nrProduse)
            totalAchizitie = 0
            totalAdaos = 0
            totalVanzare = 0
            for i in range(nrProduse):
                ws.merge_cells(start_row=11+i, start_column=2, end_row=11+i, end_column=4)
                ws['A' + str(11+i)] = i+1
                numeProdus = self.ui.tableWidget_produseLichidateClient_6.item(i,1).text()
                pretAchizitie =  int(self.ui.tableWidget_produseLichidateClient_6.item(i,4).text())
                pretVanzare =  int(self.ui.tableWidget_produseLichidateClient_6.item(i,5).text())
                ws['B' + str(11+i)] = numeProdus
                ws['E' + str(11+i)] = "B.C"
                ws['F' + str(11+i)] = 1
                ws['G' + str(11+i)] = 1
                ws['H' + str(11+i)] = pretAchizitie
                ws['I' + str(11+i)] = int(ws['G' + str(11+i)].value) * int(ws['H' + str(11+i)].value)
                ws['M' + str(11+i)] = pretVanzare
                ws['N' + str(11+i)] = pretVanzare
                adaosTotal = int(pretVanzare) - int(pretAchizitie)
                adaosProcentual =  ((int(pretVanzare)-int(pretAchizitie))/int(pretAchizitie)) *100
                adaosProcentual = round(adaosProcentual,2)
                ws['J' + str(11 + i)] = adaosProcentual
                ws['K' + str(11 + i)] = adaosTotal
                ws['L' + str(11 + i)] = adaosTotal
                totalAchizitie = totalAchizitie+pretAchizitie
                totalAdaos = totalAdaos + adaosTotal
                totalVanzare = totalVanzare + pretVanzare
                self.styleTabel(ws,11+i-1,'A')
                self.styleTabel(ws,11+i-1,'B')
                self.styleTabel(ws,11+i-1,'C')
                self.styleTabel(ws,11+i-1,'D')
                self.styleTabel(ws,11+i-1,'E')
                self.styleTabel(ws,11+i-1,'F')
                self.styleTabel(ws,11+i-1,'G')
                self.styleTabel(ws,11+i-1,'H')
                self.styleTabel(ws,11+i-1,'I')
                self.styleTabel(ws,11+i-1,'M')
                self.styleTabel(ws,11+i-1,'N')
                self.styleTabel(ws,11+i-1,'J')
                self.styleTabel(ws,11+i-1,'K')
                self.styleTabel(ws,11+i-1,'L')
            ws['G'+ str(11+nrProduse)] = nrProduse
            ws['H'+ str(11+nrProduse)] = totalAchizitie
            ws['J'+ str(11+nrProduse)] = totalAdaos
            ws['M'+ str(11+nrProduse)] = totalVanzare
            ws.merge_cells(start_row=int(11 + nrProduse), start_column=8, end_row=int(11 + nrProduse), end_column=9)
            ws.merge_cells(start_row=int(11 + nrProduse), start_column=10, end_row=int(11 + nrProduse), end_column=12)
            ws.merge_cells(start_row=int(11 + nrProduse), start_column=13, end_row=int(11 + nrProduse), end_column=14)
            for i in range(nrProduse -1,-1,-1):
                self.ui.tableWidget_produseLichidateClient_6.removeRow(i)
            self.PopulareTabelGestiuneVanzariProduseAmanet(listaDB)
            currentDir = os.getcwd()
            numeNir = currentDir + "\\NIR-uri\\"+ "A-" + str(nrNir) + ".xlsx"
            lista.save(numeNir)
            os.startfile(numeNir)
            self.ui.pushButton_goHome_25.show()
    def goToPageVanzareProduseAmanet(self):
        if not self.ui.tableWidget_produseLichidateClient_5.selectedItems():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu a fost selectat nici un contract")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            denumireProdus = self.ui.tableWidget_produseLichidateClient_5.item(self.ui.tableWidget_produseLichidateClient_5.currentRow(), 0).text()
            pretVanzare = self.ui.tableWidget_produseLichidateClient_5.item(self.ui.tableWidget_produseLichidateClient_5.currentRow(), 1).text()
            self.ui.lineEdit_NumeClient_5.setText(denumireProdus)
            self.ui.label_SerieCI_7.setText(pretVanzare)
            self.ui.stackedWidget_6.setCurrentIndex(2)
    def verificareDateIainteDeVanzareProdus(self):
        numeClient = self.ui.lineEdit_NumeClient_2.text()
        serieNr = self.ui.lineEdit_CNP_3.text()
        cnp = self.ui.lineEdit_CNP_2.text()
        adresa = self.ui.lineEdit_Adresa_2.text()
        telefon = self.ui.lineEdit_Telefon_2.text()
        email = self.ui.lineEdit_Email_2.text()
        observatii = self.ui.lineEdit_NumeClient_6.text()
        reducere = self.ui.lineEdit_PrenumeClient_5.text()
        nrBucati = self.ui.lineEdit_PrenumeClient_6.text()
        if numeClient=="" or serieNr=="" or cnp=="" or adresa=="" or telefon=="" or observatii=="" or reducere=="" or nrBucati=="":
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu au fost introduse toate datele necesare pentru a genera factura")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            return False
        else:
            return True
    def finalizareVinzareProdusAmanet(self,listaDB):
        if self.verificareDateIainteDeVanzareProdus():
            df_produse = pd.read_excel(listaDB,sheet_name="ListaProduseLichidateAmanet")
            lista = load_workbook(listaDB)
            denumireProdus = self.ui.tableWidget_produseLichidateClient_5.item(self.ui.tableWidget_produseLichidateClient_5.currentRow(), 0).text()
            observatii = self.ui.lineEdit_NumeClient_6.text()
            pretVanzare = self.ui.label_SerieCI_7.text()
            reducere = self.ui.lineEdit_PrenumeClient_5.text()
            numeClient = self.ui.lineEdit_NumeClient_2.text()
            serieNr = self.ui.lineEdit_CNP_3.text()
            cnp = self.ui.lineEdit_CNP_2.text()
            adresa = self.ui.lineEdit_Adresa_2.text()
            telefon = self.ui.lineEdit_Telefon_2.text()
            email = self.ui.lineEdit_Email_2.text()
            total = int(pretVanzare) - int(reducere)
            total = round(total,2)
            for i in range(len(df_produse.index)):
                if denumireProdus == df_produse.iloc[i]["Denumire Produs"]:
                    ws = lista["ListaProduseLichidateAmanet"]
                    ws.delete_rows(i + 2)
                    ws = lista["ListaProduseLichidateAm-Vandute"]
                    row = df_produse.loc[i].values.tolist()
                    row.append(int(pretVanzare)-int(reducere))
                    row[16] = date.today().strftime("%d.%m.%Y")
                    ws.append(row)
            denumireProdus = self.ui.lineEdit_NumeClient_5.text()
            nrFactura = self.genereazaNrFactura(listaDB)
            lista.save(listaDB)
            ws = lista["templateFactura"]
            ws["E5"] = "Seria EPSA - NR."+ str(nrFactura)
            ws["E6"] = "DATA: "+ date.today().strftime("%d.%m.%Y")
            ws["J1"] = numeClient
            ws["J2"] = adresa
            ws["J3"] = serieNr
            ws["J4"] = cnp
            ws["J5"] = telefon
            ws["J6"] = email
            ws["A12"] = "1"
            ws["B12"] = denumireProdus + " "+observatii
            ws["H12"] = "B.C"
            ws["I12"] = "1"
            ws["J12"] = round(int(pretVanzare),2)
            ws["K12"] = round(int(pretVanzare),2)
            ws["J35"] = round(int(pretVanzare),2)
            ws["J36"] = round(int(reducere),2)
            ws["J37"] = total
            ws["A39"] = "DATA: " + date.today().strftime("%d.%m.%Y")
            currentDir = os.getcwd()
            numeFactura = currentDir + "\\Facturi\\" + "A-" + str(nrFactura) + ".xlsx"
            lista.save(numeFactura)
            os.startfile(numeFactura)
            self.ui.stackedWidget.setCurrentIndex(0)
    def genereazaNrFactura(self,listaDB):
        lista = load_workbook(listaDB)
        ws = lista["ListaFacturi"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        ws['B' + str(maxRow + 1)] = date.today().strftime("%d.%m.%Y")
        self.styleTabel(ws, maxRow, 'A')
        self.styleTabel(ws, maxRow, 'B')
        lista.save(listaDB)
        return maxRow
    def prelungesteContract(self,listaDB):
        print("START prelungesteContract")
        tabelCurent = self.ui.stackedWidget_4.currentIndex()
        listaTabele = [self.ui.tableWidget_contracteInTermen,self.ui.tableWidget_contracteInAsteptare,self.ui.tableWidget_contracteExpirate,self.ui.tableWidget_contracteIncheiate,self.ui.tableWidget_contractePrelungite]
        listaContracte = ["ListaContracteInTermen","ListaContracteInAsteptare","ListaContracteExpirate","ListaContracteLichidate","ListaContractePrelungite"]
        listaProduse = ["ListaProduseInTermen","ListaProduseInAsteptare","ListaProduseExpirate","ListaProduseLichidateClient","ListaProdusePrelungite"]
        self.ui.frame_31.show()
        if not listaTabele[tabelCurent].selectedItems():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu a fost selectat nici un contract")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            indexRand = None
            randContract = None
            for i in range(len(listaTabele)):
                if listaTabele[i].selectedItems():
                    tabelCurent = i
                    indexRand = listaTabele[tabelCurent].selectedIndexes()
                    randContract = indexRand[0].row()
                    break
            nrContract = listaTabele[tabelCurent].item(listaTabele[tabelCurent].currentRow(), 0).text()
            # Preia datele contractului din DB
            df_contract = pd.read_excel(listaDB, sheet_name=listaContracte[tabelCurent])
            for i in range(len(df_contract.index)):
                if int(nrContract) == df_contract.iloc[i]["Nr. Contract"]:
                    df_contract = df_contract.loc[i]
                    break
            # Preia date produse din contractul selectat
            df_toateProdusele = pd.read_excel(listaDB,sheet_name=listaProduse[tabelCurent]) # Iau toate produsele
            df_produse = pd.DataFrame()
            for i in range(len(df_toateProdusele.index)):
                if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                    df_produse = df_produse.append(df_toateProdusele.loc[i])
            # Preia date client dupa nr contractului
            df_client = pd.read_excel(listaDB, sheet_name="ListaClienti")
            for i in range(len(df_client.index)):
                print(df_client.iloc[i]["Contracte "],type(df_client.iloc[i]["Contracte "]))
                contract = df_client.iloc[i]["Contracte "]
                if isinstance(contract,int): # daca clientul are un singur contract in db
                    if df_client.iloc[i]["Contracte "] == int(nrContract):
                        df_client = df_client.loc[i]
                        break
                else: # clientul are mai multe contracte in db
                    contract = contract.replace(" ","")
                    contract= contract.split(",")
                    print(contract)
                    if str(nrContract) in contract:
                        df_client = df_client.loc[i]
                        break
            if tabelCurent == 0: # Daca contractul este In Termen
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " este in termen\nSunteti sigur ca doriti sa il prelungiti?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.hide()
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    # comision = self.calculeazaComisionPentruPrelungireaContractelorPrelungiteAnterior(df_contract)
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_22.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_15.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_14.setText("0")
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.show()
                    self.ui.pushButton_finalizareIncheiereContract.hide()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost prelungit!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
            elif tabelCurent == 1: # Daca contractul este in asteptare
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " este in perioada de asteptare\nSe va percepe comision de penalizare\nSunteti sigur ca doriti sa il prelungiti?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.hide()
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_22.setText(str(df_contract["Comision"]))
                    penalizari = self.calculeazaPenalizari(df_contract)
                    self.ui.label_numarCercDreapta_14.setText(str(penalizari)) #penalizari
                    totalDePlata = penalizari + df_contract["Comision"]
                    self.ui.label_numarCercDreapta_15.setText(str(totalDePlata))  # Total comision curent + penalizari
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.show()
                    self.ui.pushButton_finalizareIncheiereContract.hide()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost prelungit!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
            elif tabelCurent == 2:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " nu se poate prelungi pentru ca este EXPIRAT\n")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec_()
                if returnValue == QMessageBox.Ok:
                    print("OK clicked")
            elif tabelCurent == 3:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " nu se poate prelungi pentru ca este INCHEIAT\n")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec_()
                if returnValue == QMessageBox.Ok:
                    print("OK clicked")
            # if pd.isnull(df_contract["Data expirare initiala"]): # Verific daca contractul a mai fost prelungit
            #     # Daca NU:
            #     print("aici########")
            elif tabelCurent == 4: # Daca contractul este prelungit
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " a mai fost prelungit\nSunteti sigur ca doriti sa il prelungiti din nou?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.show()
                    ultimaDataDePrelungire = df_contract["Data prelungire"].split(",")
                    ultimaDataDePrelungire = ultimaDataDePrelungire[0]
                    nrZilePrelungitUltimaData = str(df_contract["Nr. Zile prelungite"]).split(",")
                    nrZilePrelungitUltimaData = nrZilePrelungitUltimaData[0]
                    self.ui.label_numarCercDreapta_6.setText(ultimaDataDePrelungire)# ultima data de prelungire
                    self.ui.label_numarCercDreapta_5.setText(df_contract["Data expirare"])# noua data de expirare
                    self.ui.label_textSusCercMijloc_46.setText(nrZilePrelungitUltimaData+" zile")# nr de zile cu care a fost prelungit ultima data
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    comisionCurent = str(df_contract["Comision prelungire"]).split(",")
                    comisionCurent = comisionCurent[0]
                    self.ui.label_numarCercDreapta_22.setText(str(comisionCurent))
                    # penalizari = self.calculeazaPenalizari(df_contract)
                    self.ui.label_numarCercDreapta_14.setText("0") #penalizari
                    # totalDePlata = penalizari + df_contract["Comision"]
                    self.ui.label_numarCercDreapta_15.setText(str(comisionCurent))
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.show()
                    self.ui.pushButton_finalizareIncheiereContract.hide()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost prelungit!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
        print("STOP prelungesteContract- no return:")
    def incheieContract(self,listaDB):
        print("START incheieContract")
        tabelCurent = self.ui.stackedWidget_4.currentIndex()
        listaTabele = [self.ui.tableWidget_contracteInTermen, self.ui.tableWidget_contracteInAsteptare,
                       self.ui.tableWidget_contracteExpirate, self.ui.tableWidget_contracteIncheiate,
                       self.ui.tableWidget_contractePrelungite]
        listaContracte = ["ListaContracteInTermen", "ListaContracteInAsteptare", "ListaContracteExpirate",
                          "ListaContracteLichidate", "ListaContractePrelungite"]
        listaProduse = ["ListaProduseInTermen", "ListaProduseInAsteptare", "ListaProduseExpirate",
                        "ListaProduseLichidateClient", "ListaProdusePrelungite"]
        self.ui.frame_31.hide()
        if not listaTabele[tabelCurent].selectedItems():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu a fost selectat nici un contract")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
        else:
            indexRand = None
            randContract = None
            for i in range(len(listaTabele)):
                if listaTabele[i].selectedItems():
                    tabelCurent = i
                    indexRand = listaTabele[tabelCurent].selectedIndexes()
                    randContract = indexRand[0].row()
                    break
            nrContract = listaTabele[tabelCurent].item(listaTabele[tabelCurent].currentRow(), 0).text()
            # Preia datele contractului din DB
            df_contract = pd.read_excel(listaDB, sheet_name=listaContracte[tabelCurent])
            for i in range(len(df_contract.index)):
                if int(nrContract) == df_contract.iloc[i]["Nr. Contract"]:
                    df_contract = df_contract.loc[i]
                    break
            # Preia date produse din contractul selectat
            df_toateProdusele = pd.read_excel(listaDB, sheet_name=listaProduse[tabelCurent])  # Iau toate produsele
            df_produse = pd.DataFrame()
            for i in range(len(df_toateProdusele.index)):
                if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                    df_produse = df_produse.append(df_toateProdusele.loc[i])
            # Preia date client dupa nr contractului
            df_client = pd.read_excel(listaDB, sheet_name="ListaClienti")
            for i in range(len(df_client.index)):
                print(df_client.iloc[i]["Contracte "], type(df_client.iloc[i]["Contracte "]))
                contract = df_client.iloc[i]["Contracte "]
                if isinstance(contract, int):  # daca clientul are un singur contract in db
                    if df_client.iloc[i]["Contracte "] == int(nrContract):
                        df_client = df_client.loc[i]
                        break
                else:  # clientul are mai multe contracte in db
                    contract = contract.replace(" ", "")
                    contract = contract.split(",")
                    print(contract)
                    if str(nrContract) in contract:
                        df_client = df_client.loc[i]
                        break
            if tabelCurent == 0:  # Daca contractul este In Termen
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " este in termen\nSunteti sigur ca doriti sa il lichidati?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.hide()
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    # Calculez comisionul de achitat in functie de cate zile au trecut de la crearea contractului
                    nrZileScurse = date.today() - datetime.strptime(str(df_contract["Data intrare in amanet"]),
                                                                    '%d.%m.%Y').date()
                    nrZileScurse = nrZileScurse.days
                    if nrZileScurse <=10:
                        nrZileScurse = 10
                    comisionCurent = (int(df_contract["Comision"])/30) * nrZileScurse
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_22.setText(str(comisionCurent))
                    # Total = suma imprumutata + comision curent
                    total = int(df_contract["Suma imprumutata"]) + comisionCurent
                    self.ui.label_numarCercDreapta_15.setText(str(total))
                    self.ui.label_numarCercDreapta_14.setText("0")
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.hide()
                    self.ui.pushButton_finalizareIncheiereContract.show()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost lichidat!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
            elif tabelCurent == 1:  # Daca contractul este in asteptare
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " este in perioada de asteptare\nSe va percepe comision de penalizare\nSunteti sigur ca doriti sa il lichidati?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.hide()
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    self.ui.label_numarCercDreapta_22.setText(str(df_contract["Comision"]))
                    penalizari = self.calculeazaPenalizari(df_contract)
                    self.ui.label_numarCercDreapta_14.setText(str(penalizari))  # penalizari
                    totalDePlata = penalizari + df_contract["Comision"] + int(df_contract["Suma imprumutata"])
                    self.ui.label_numarCercDreapta_15.setText(str(totalDePlata))  # Total comision curent + penalizari
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.hide()
                    self.ui.pushButton_finalizareIncheiereContract.show()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost lichidat!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
            elif tabelCurent == 2:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " nu se poate prelungi pentru ca este EXPIRAT\n")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec_()
                if returnValue == QMessageBox.Ok:
                    print("OK clicked")
            elif tabelCurent == 3:
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " nu se poate prelungi pentru ca este INCHEIAT\n")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec_()
                if returnValue == QMessageBox.Ok:
                    print("OK clicked")
            # if pd.isnull(df_contract["Data expirare initiala"]): # Verific daca contractul a mai fost prelungit
            #     # Daca NU:
            #     print("aici########")
            elif tabelCurent == 4:  # Daca contractul este prelungit
                msgBox = QMessageBox()
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText(
                    "Contractul " + nrContract + " a mai fost prelungit\nSunteti sigur ca doriti sa il lichidati?")
                msgBox.setWindowTitle("Atentie")
                msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                returnValue = msgBox.exec()
                if returnValue == QMessageBox.Yes:
                    produse = ""
                    observatiiProduse = ""
                    for j in range(len(df_produse.index)):
                        produse = produse + df_produse.iloc[j]["Denumire Produs"] + "&"
                        observatiiProduse = observatiiProduse + df_produse.iloc[j]["Observatii"] + "&"
                    produse = produse[:-1]
                    observatiiProduse = observatiiProduse[:-1]
                    self.ui.frame_66.show()
                    ultimaDataDePrelungire = df_contract["Data prelungire"].split(",")
                    ultimaDataDePrelungire = ultimaDataDePrelungire[0]
                    nrZilePrelungitUltimaData = str(df_contract["Nr. Zile prelungite"]).split(",")
                    nrZilePrelungitUltimaData = nrZilePrelungitUltimaData[0]
                    self.ui.label_numarCercDreapta_6.setText(ultimaDataDePrelungire)  # ultima data de prelungire
                    self.ui.label_numarCercDreapta_5.setText(df_contract["Data expirare"])  # noua data de expirare
                    self.ui.label_textSusCercMijloc_46.setText(
                        nrZilePrelungitUltimaData + " zile")  # nr de zile cu care a fost prelungit ultima data
                    self.ui.label_numarCercDreapta_3.setText(str(df_contract["Data intrare in amanet"]))
                    self.ui.label_numarCercDreapta_4.setText(str(df_contract["Nume"]))
                    self.ui.label_numarCercDreapta_17.setText(str(df_contract["Prenume"]))
                    self.ui.label_numarCercDreapta_20.setText(str(df_client["CNP"]))
                    self.ui.label_numarCercDreapta_19.setText(str(df_client["Serie C.I"]))
                    self.ui.label_numarCercDreapta_16.setText(str(df_client["Numar C.I"]))
                    self.ui.label_numarCercDreapta_18.setText(str(df_client["Telefon"]))
                    self.ui.label_numarCercDreapta_11.setText(str(df_client["Email"]))
                    self.ui.label_numarCercMijloc_9.setText(produse)
                    self.ui.label_numarCercMijloc_10.setText(observatiiProduse)
                    self.ui.label_numarCercDreapta_12.setText(str(df_contract["Suma imprumutata"]))
                    self.ui.label_numarCercDreapta_13.setText(str(df_contract["Comision"]))
                    df_notaPrelungire = pd.read_excel(listaDB, sheet_name="ListaNotePrelungire")
                    for i in range(len(df_notaPrelungire.index)-1,-1,-1):
                        print(df_notaPrelungire.iloc[i]["Nr. Contract prelungit"])
                        if int(nrContract) == int(df_notaPrelungire.iloc[i]["Nr. Contract prelungit"]):
                            dataExpirare = df_notaPrelungire.iloc[i]["Data prelungire"] # aici modific daca se calculeaza de la data prelungirii sau de la data expirarii anterioare
                            print(dataExpirare)
                            break
                    nrZileScurse = date.today() - datetime.strptime(dataExpirare,'%d.%m.%Y').date()
                    nrZileScurse = nrZileScurse.days
                    if nrZileScurse <=0 :
                        print("AICI zile scurse <=0")
                        nrZileScurse =0
                    comisionCurent = (int(df_contract["Comision"]) / 30) * nrZileScurse
                    self.ui.label_numarCercDreapta_22.setText(str(comisionCurent))
                    self.ui.label_numarCercDreapta_14.setText("0")  # penalizari
                    totalDePlata = comisionCurent +int(df_contract["Suma imprumutata"])
                    self.ui.label_numarCercDreapta_15.setText(str(totalDePlata))
                    self.ui.stackedWidget_4.setCurrentIndex(5)
                    self.ui.label_numarCercDreapta_21.hide()
                    self.ui.label_textSusCercMijloc_40.hide()
                    self.ui.pushButton_incheieContract_7.hide()
                    self.ui.pushButton_incheieContract_6.hide()
                    self.ui.pushButton_incheieContract_8.hide()
                    self.ui.pushButton_finalizareIncheiereContract.show()
                    self.ui.label_numarCercDreapta_2.setText(nrContract)
                elif returnValue == QMessageBox.No:
                    msgBox1 = QMessageBox()
                    msgBox1.setIcon(QMessageBox.Information)
                    msgBox1.setText(
                        "Contractul " + nrContract + " NU a fost prelungit!")
                    msgBox1.setWindowTitle("Atentie")
                    msgBox1.setStandardButtons(QMessageBox.Ok)
                    msgBox1.exec()
        print("STOP incheieContract- no return:")

    def calculeazaPenalizari(self,df_contract):
        nrZileIntarziate = date.today() - datetime.strptime(df_contract["Data expirare"], '%d.%m.%Y').date()
        nrZileIntarziate = nrZileIntarziate.days
        print(nrZileIntarziate,"NR ZILE INTARZIATE")
        # mai jos 0.5 reprezinta 0.5%/zi pentru zilele scurse dupa data de expirare
        penalizari = (0.5 / 100) * int(df_contract["Suma imprumutata"]) * nrZileIntarziate
        print(penalizari)
        penalizari = round(penalizari, 2)
        return penalizari
    def calculeazaComisionDupaXZilePrelungire(self):
        print("START calculeazaComisionDupaXZilePrelungire")
        nrZilePrelungire = self.ui.lineEdit_numarContractDePrelungit_7.text()
        if nrZilePrelungire == "" or not nrZilePrelungire.isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introdu numarul de zile pentru care doresti sa prelungesti contractul")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            msgBox.exec()
        else:
            text = "Comisionul de plata dupa " + nrZilePrelungire + " zile va fi:"
            #calculeaza comisionul dupa x zile prelungite
            comision = self.ui.label_numarCercDreapta_13.text()
            comision = int(comision)
            comision = (comision/30) * int(nrZilePrelungire)
            comision = round(comision,2)
            self.ui.label_textSusCercMijloc_40.setText(text )
            self.ui.label_numarCercDreapta_21.setText(str(comision))
            self.ui.label_numarCercDreapta_21.show()
            self.ui.label_textSusCercMijloc_40.show()
            print("STOP calculeazaComisionDupaXZilePrelungire - return comision:", comision)
            return comision
    def finalizarePrelungireContract(self,listaDB):
        print("START finalizarePrelungireContract")
        nrContract = self.ui.label_numarCercDreapta_2.text()
        nrZilePrelungire = self.ui.lineEdit_numarContractDePrelungit_7.text()
        tabelCurent = self.ui.stackedWidget_4.currentIndex()
        listaTabele = [self.ui.tableWidget_contracteInTermen, self.ui.tableWidget_contracteInAsteptare,
                       self.ui.tableWidget_contracteExpirate, self.ui.tableWidget_contracteIncheiate,
                       self.ui.tableWidget_contractePrelungite]
        listaContracte = ["ListaContracteInTermen", "ListaContracteInAsteptare", "ListaContracteExpirate",
                          "ListaContracteLichidate", "ListaContractePrelungite"]
        listaProduse = ["ListaProduseInTermen", "ListaProduseInAsteptare", "ListaProduseExpirate",
                        "ListaProduseLichidateClient", "ListaProdusePrelungite"]
        if self.ui.lineEdit_numarContractDePrelungit_6.text() == "":
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introdu numarul bonului fiscal")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            msgBox.exec()
        elif nrZilePrelungire == "" or not nrZilePrelungire.isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introdu numarul de zile pentru care doresti sa prelungesti contractul")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            msgBox.exec()
        else:
            indexRand = None
            randContract = None
            for i in range(len(listaTabele)):
                if listaTabele[i].selectedItems():
                    tabelCurent = i
                    indexRand = listaTabele[tabelCurent].selectedIndexes()
                    randContract = indexRand[0].row()
                    break
            nrContract = listaTabele[tabelCurent].item(listaTabele[tabelCurent].currentRow(), 0).text()
            # Preia datele contractului din DB
            df_contract = pd.read_excel(listaDB, sheet_name=listaContracte[tabelCurent])
            randContract = 0
            for i in range(len(df_contract.index)):
                if int(nrContract) == df_contract.iloc[i]["Nr. Contract"]:
                    df_contract = df_contract.loc[i]
                    randContract = i+2
                    break
            # Preia date produse din contractul selectat
            df_toateProdusele = pd.read_excel(listaDB, sheet_name=listaProduse[tabelCurent])  # Iau toate produsele
            df_produse = pd.DataFrame(columns=["Denumire Produs","Observatii","Valoare Evaluata","Suma Imprumutata","Comision(%/zi)","Comision (lei)","Data intrare in amanet","Data expirare","Total de restituit (lei)","Nr. Contract"])
            randuriProduseDeSters = []
            for i in range(len(df_toateProdusele.index)):
                if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                    df_produse = df_produse.append(df_toateProdusele.loc[i])
                    randuriProduseDeSters.append(i+2)
            # Preia date client dupa nr contractului
            df_client = pd.read_excel(listaDB, sheet_name="ListaClienti")
            for i in range(len(df_client.index)):
                print(df_client.iloc[i]["Contracte "], type(df_client.iloc[i]["Contracte "]))
                contract = df_client.iloc[i]["Contracte "]
                if isinstance(contract, int):  # daca clientul are un singur contract in db
                    if df_client.iloc[i]["Contracte "] == int(nrContract):
                        df_client = df_client.loc[i]
                        break
                else:  # clientul are mai multe contracte in db
                    contract = contract.replace(" ", "")
                    contract = contract.split(",")
                    print(contract)
                    if str(nrContract) in contract:
                        df_client = df_client.loc[i]
                        break
            lista = load_workbook(listaDB)
            row = df_contract.values.tolist()
            dataExpirare = datetime.strptime(row[6], '%d.%m.%Y').date()
            dataPrelungita = (dataExpirare + timedelta(days=int(nrZilePrelungire))).strftime('%d.%m.%Y')
            if tabelCurent == 0: # Contractul este in termen
                # adaug Contractul in lista cu contracte prelungite
                row.append(row[6])
                row[6] = dataPrelungita
                row.append(date.today().strftime('%d.%m.%Y'))
                row.append(nrZilePrelungire)
                comisionDupaPrelungire = self.calculeazaComisionDupaXZilePrelungire()
                row.append(comisionDupaPrelungire)
                totalDupaPrelungire = row[7] + comisionDupaPrelungire
                row.append(totalDupaPrelungire)
                ws = lista["ListaContractePrelungite"]
                ws.append(row)
                # Mut contractul si in aplicatie
                # Din tabelul cu contracte in termen, in tabelul cu contracte prelungite
                nrRanduri = self.ui.tableWidget_contractePrelungite.rowCount()
                self.ui.tableWidget_contractePrelungite.setRowCount(nrRanduri+1)
                for i in range(len(row)):
                    self.ui.tableWidget_contractePrelungite.setItem(nrRanduri,i,QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contracteInTermen.removeRow(self.ui.tableWidget_contracteInTermen.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws = lista["ListaContracteInTermen"]
                ws.delete_rows(randContract)
                # Fac aceleasi mutari si pentru produsele din contract
                # Adaug produsele in lista cu produse prelungite
                row = df_produse.values.tolist()
                dataExpirare = datetime.strptime(row[0][7], '%d.%m.%Y').date()
                dataPrelungita = (dataExpirare + timedelta(days=int(nrZilePrelungire))).strftime('%d.%m.%Y')
                for i in range(len(row)-1,-1,-1):
                    print(row[i])
                    row[i].append(row[i][7])
                    row[i][7] = dataPrelungita
                    row[i].append(date.today().strftime('%d.%m.%Y'))
                    row[i].append(nrZilePrelungire)
                    comisionDupaPrelungire = (row[i][5]/30) * int(nrZilePrelungire)
                    comisionDupaPrelungire = round(comisionDupaPrelungire,2)
                    row[i].append(comisionDupaPrelungire)
                    totalDupaPrelungire = row[i][3] + comisionDupaPrelungire
                    row[i].append(totalDupaPrelungire)
                    ws = lista["ListaProdusePrelungite"]
                    ws.append(row[i])
                    print(row[i])
                    ws = lista["ListaProduseInTermen"]
                    ws.delete_rows(randuriProduseDeSters[i])
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaPrelungire(lista, listaDB, nrContract,
                                             dataExpirare,dataPrelungita)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                    "Contractul " + nrContract + " a fost prelungit!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
            elif tabelCurent == 1: # Daca contractul este in asteptare
                # adaug Contractul in lista cu contracte prelungite
                row = df_contract.values.tolist()
                print(row)
                dataExpirare = datetime.strptime(row[6], '%d.%m.%Y').date()
                dataPrelungita = (dataExpirare + timedelta(days=int(nrZilePrelungire))).strftime('%d.%m.%Y')
                row.append(row[6])
                row[6] = dataPrelungita
                row.append(date.today().strftime('%d.%m.%Y'))
                row.append(nrZilePrelungire)
                comisionDupaPrelungire = self.calculeazaComisionDupaXZilePrelungire()
                row.append(comisionDupaPrelungire)
                totalDupaPrelungire = row[7] + comisionDupaPrelungire
                row.append(totalDupaPrelungire)
                ws = lista["ListaContractePrelungite"]
                ws.append(row)
                # Mut contractul si in aplicatie
                # Din tabelul cu contracte in asteptare, in tabelul cu contracte prelungite
                nrRanduri = self.ui.tableWidget_contractePrelungite.rowCount()
                self.ui.tableWidget_contractePrelungite.setRowCount(nrRanduri + 1)
                for i in range(len(row)):
                    self.ui.tableWidget_contractePrelungite.setItem(nrRanduri, i,
                                                                    QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contracteInAsteptare.removeRow(self.ui.tableWidget_contracteInAsteptare.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws = lista["ListaContracteInAsteptare"]
                ws.delete_rows(randContract)
                # Fac aceleasi mutari si pentru produsele din contract
                # Adaug produsele in lista cu produse prelungite
                row = df_produse.values.tolist()
                for i in range(len(row) - 1, -1, -1):
                    dataExpirare = datetime.strptime(row[i][7], '%d.%m.%Y').date()
                    dataPrelungita = (dataExpirare + timedelta(days=int(nrZilePrelungire))).strftime('%d.%m.%Y')
                    row[i].append(row[i][7])
                    row[i][7] = dataPrelungita
                    row[i].append(date.today().strftime('%d.%m.%Y'))
                    row[i].append(nrZilePrelungire)
                    comisionDupaPrelungire = (row[i][5] / 30) * int(nrZilePrelungire)
                    comisionDupaPrelungire = round(comisionDupaPrelungire, 2)
                    row[i].append(comisionDupaPrelungire)
                    totalDupaPrelungire = row[i][3] + comisionDupaPrelungire
                    row[i].append(totalDupaPrelungire)
                    ws = lista["ListaProdusePrelungite"]
                    ws.append(row[i])
                    ws = lista["ListaProduseInAsteptare"]
                    ws.delete_rows(randuriProduseDeSters[i])
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaPrelungire(lista, listaDB, nrContract,
                                                 dataExpirare, dataPrelungita)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                        "Contractul " + nrContract + " a fost prelungit!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
            elif tabelCurent == 4:  # Daca contractul a mai fost prelungit
                # Fac update in lista de contracte prelungite pentru contractul prelungit
                ws = lista["ListaContractePrelungite"]
                row = df_contract.values.tolist()
                dataExpirare = datetime.strptime(row[6], '%d.%m.%Y').date()
                dataPrelungita = (dataExpirare + timedelta(days=int(nrZilePrelungire))).strftime('%d.%m.%Y')
                row[6] = dataPrelungita
                row[11] = date.today().strftime('%d.%m.%Y') + ", " + row[11]
                row[12] = nrZilePrelungire + ", " + str(row[12])
                comisionDupaPrelungire = self.calculeazaComisionDupaXZilePrelungire()
                row[13] = str(comisionDupaPrelungire) + ", " + str(row[13])
                totalDupaPrelungire = row[7] + comisionDupaPrelungire
                row[14] = str(totalDupaPrelungire) + ", " + str(row[14])
                ws.append(row)
                # Fac update si in tabelul din aplicatie
                # Din tabelul cu contracte in asteptare, in tabelul cu contracte prelungite
                nrRanduri = self.ui.tableWidget_contractePrelungite.rowCount()
                self.ui.tableWidget_contractePrelungite.setRowCount(nrRanduri + 1)
                for i in range(len(row)):
                    self.ui.tableWidget_contractePrelungite.setItem(nrRanduri, i,
                                                                    QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contractePrelungite.removeRow(self.ui.tableWidget_contractePrelungite.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws.delete_rows(randContract)
                #preiau date produs pentru contractele prelungite deja
                df_toateProdusele = pd.read_excel(listaDB, sheet_name=listaProduse[tabelCurent])  # Iau toate produsele
                df_produse = pd.DataFrame(
                    columns=["Denumire Produs", "Observatii", "Valoare Evaluata", "Suma Imprumutata", "Comision(%/zi)",
                             "Comision (lei)", "Data intrare in amanet", "Data expirare", "Total de restituit (lei)",
                             "Nr. Contract","Data expirare initiala","Data prelungire","Nr. Zile prelungite","Comision prelungire","Total de restituit dupa prelungire"])
                randuriProduseDeSters = []
                for i in range(len(df_toateProdusele.index)):
                    if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                        df_produse = df_produse.append(df_toateProdusele.loc[i])
                        randuriProduseDeSters.append(i + 2)
                # Fac aceleasi mutari si pentru produsele din contract
                # Update in lista cu produse prelungite
                row = df_produse.values.tolist()
                for i in range(len(row) - 1, -1, -1):
                    print(row[i])
                    row[i][7] = dataPrelungita
                    row[i][11] =  date.today().strftime('%d.%m.%Y') + ", " + row[i][11]
                    row[i][12] = nrZilePrelungire + ", " + str(row[i][12])
                    comisionDupaPrelungire = (int(row[i][5])/30) * int(nrZilePrelungire)
                    row[i][13] = str(comisionDupaPrelungire) + ", " + str(row[i][13])
                    totalDupaPrelungire = row[i][3] + comisionDupaPrelungire
                    row[i][14] = str(totalDupaPrelungire) + ", " + str(row[i][14])
                    ws = lista["ListaProdusePrelungite"]
                    ws.append(row[i])
                    print(row[i])
                    ws.delete_rows(randuriProduseDeSters[i])
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaPrelungire(lista, listaDB, nrContract,
                                             dataExpirare, dataPrelungita)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                    "Contractul " + nrContract + " a fost prelungit!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
        print("STOP finalizarePrelungireContract- no return:")
    def genereazaNumarNotaPrelungire(self,lista,listaDB,nrContract,dataExpirare,nrZilePrelungite):
        print("START genereazaNumarContract")
        ws = lista["ListaNotePrelungire"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        ws['B' + str(maxRow + 1)] = date.today().strftime("%d.%m.%Y")
        ws['C' + str(maxRow + 1)] = nrContract
        ws['D' + str(maxRow + 1)] = dataExpirare.strftime("%d.%m.%Y")
        ws['E' + str(maxRow + 1)] = nrZilePrelungite
        self.styleTabel(ws, maxRow, 'A')
        self.styleTabel(ws, maxRow, 'B')
        self.styleTabel(ws, maxRow, 'C')
        self.styleTabel(ws, maxRow, 'D')
        self.styleTabel(ws, maxRow, 'E')
        lista.save(listaDB)
        print("STOP genereazaNumarContract- return:", maxRow)
        return maxRow
    def genereazaNotaPrelungire(self, lista, listaDB, nrContract, dataExpirare,dataRestituire):
        print("START genereazaNotaPrelungire")
        currentDir = os.getcwd()
        # Genereaza nr nota de prelungire
        nrZilePrelungite = self.ui.lineEdit_numarContractDePrelungit_7.text()
        numarNotaPrelungire = self.genereazaNumarNotaPrelungire(lista, listaDB, nrContract, dataExpirare,
                                                                nrZilePrelungite)
        # dataExpirare = datetime.strptime(dataExpirare, "%d.%m.%Y").date()
        # dataRestituire = (dataExpirare + timedelta(days=int(nrZilePrelungite))).strftime('%d.%m.%Y')
        # 2.Completeaza NOTA DE Prelungire
        template = "/templateContract/TemplateNotaDePrelungire.docx"
        copyTemplate = "/notePrelungire"
        originalTempl = currentDir + template
        copyTempl = currentDir + copyTemplate
        shutil.copy(originalTempl, copyTempl)
        document = docx.Document(currentDir + "/notePrelungire/TemplateNotaDePrelungire.docx")
        today = date.today().strftime("%d.%m.%Y")
        numeClient = self.ui.label_numarCercDreapta_4.text() + " " + self.ui.label_numarCercDreapta_17.text()
        numeClient = numeClient.replace(" ","-")
        numeContract = str(nrContract) + "-" + numeClient+ ".docx"
        numeNotaPrelungire = str(numarNotaPrelungire) + "-" + nrContract + "-" + numeClient + ".docx"
        numeClient = numeClient.replace("-", " ")
        cnp = self.ui.label_numarCercDreapta_20.text()
        sumaImprumutata = self.ui.label_numarCercDreapta_12.text()
        comision = self.ui.label_numarCercDreapta_22.text()
        penalizari = self.ui.label_numarCercDreapta_14.text()
        total = self.ui.label_numarCercDreapta_15.text()
        nrBon = self.ui.lineEdit_numarContractDePrelungit_6.text()
        dataCreareContract = self.ui.label_numarCercDreapta_3.text()
        comisionLaFinal = self.calculeazaComisionDupaXZilePrelungire()
        totalLaFinal = int(sumaImprumutata) + comisionLaFinal
        totalLaFinal = round(totalLaFinal,2)
        for paragraph in document.paragraphs:
            runs = paragraph.runs
            for i in range(len(runs)):
                if "[NR_PRELUNGIRE]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_PRELUNGIRE]", str(numarNotaPrelungire))
                    runs[i].bold = True
                if "[DATA]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[DATA]", today)
                    runs[i].bold = True
                if "[NR_CONTRACT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_CONTRACT]", str(nrContract))
                    runs[i].bold = True
                if "[DATA_CONTRACT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[DATA_CONTRACT]", dataCreareContract)
                    runs[i].bold = True
                if "[NUME_CLIENT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NUME_CLIENT]", numeClient)
                    runs[i].bold = True
                if "[CNP_CLIENT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[CNP_CLIENT]", cnp)
                    runs[i].bold = True
                if "[COMISION]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[COMISION]", comision)
                    runs[i].bold = True
                if "[PENALIZARI]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[PENALIZARI]", penalizari)
                    runs[i].bold = True
                if "[TOTAL]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[TOTAL]", total)
                    runs[i].bold = True
                if "[NR_BON]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_BON]", nrBon)
                    runs[i].bold = True
                if "[NR_ZILE_PRELUNGITE]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_ZILE_PRELUNGITE]", nrZilePrelungite)
                    runs[i].bold = True
                if "[COMISION_LA_FINAL]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[COMISION_LA_FINAL]", str(comisionLaFinal))
                    runs[i].bold = True
                if "[TOTAL_LA_FINAL]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[TOTAL_LA_FINAL]", str(totalLaFinal))
                    runs[i].bold = True
                if "[DATA_RESTITUIRE]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[DATA_RESTITUIRE]", dataRestituire)
                    runs[i].bold = True
        document.save(currentDir + "/notePrelungire/" + numeNotaPrelungire)
        os.system('open '+currentDir + "/notePrelungire/" + numeNotaPrelungire)
        #os.startfile(currentDir + "\\notePrelungire\\" + numeNotaPrelungire)
        print("STOP genereazaNotaPrelungire - return: notaDePrelungire.docx")
    def finalizareIncheiereContract(self,listaDB):
        print("START finalizareIncheiereContract")
        nrContract = self.ui.label_numarCercDreapta_2.text()
        nrZilePrelungire = self.ui.lineEdit_numarContractDePrelungit_7.text()
        tabelCurent = self.ui.stackedWidget_4.currentIndex()
        listaTabele = [self.ui.tableWidget_contracteInTermen, self.ui.tableWidget_contracteInAsteptare,
                       self.ui.tableWidget_contracteExpirate, self.ui.tableWidget_contracteIncheiate,
                       self.ui.tableWidget_contractePrelungite]
        listaContracte = ["ListaContracteInTermen", "ListaContracteInAsteptare", "ListaContracteExpirate",
                          "ListaContracteLichidate", "ListaContractePrelungite"]
        listaProduse = ["ListaProduseInTermen", "ListaProduseInAsteptare", "ListaProduseExpirate",
                        "ListaProduseLichidateClient", "ListaProdusePrelungite"]
        if self.ui.lineEdit_numarContractDePrelungit_6.text() == "":
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introdu numarul bonului fiscal")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            msgBox.exec()
        else:
            indexRand = None
            randContract = None
            for i in range(len(listaTabele)):
                if listaTabele[i].selectedItems():
                    tabelCurent = i
                    indexRand = listaTabele[tabelCurent].selectedIndexes()
                    randContract = indexRand[0].row()
                    break
            nrContract = listaTabele[tabelCurent].item(listaTabele[tabelCurent].currentRow(), 0).text()
            # Preia datele contractului din DB
            df_contract = pd.read_excel(listaDB, sheet_name=listaContracte[tabelCurent])
            randContract = 0
            for i in range(len(df_contract.index)):
                if int(nrContract) == df_contract.iloc[i]["Nr. Contract"]:
                    df_contract = df_contract.loc[i]
                    randContract = i + 2
                    break
            # Preia date produse din contractul selectat
            df_toateProdusele = pd.read_excel(listaDB, sheet_name=listaProduse[tabelCurent])  # Iau toate produsele
            df_produse = pd.DataFrame(
                columns=["Denumire Produs", "Observatii", "Valoare Evaluata", "Suma Imprumutata", "Comision(%/zi)",
                         "Comision (lei)", "Data intrare in amanet", "Data expirare", "Total de restituit (lei)",
                         "Nr. Contract", "Data expirare initiala", "Data prelungire", "Nr. Zile prelungite",
                         "Comision prelungire", "Total de restituit dupa prelungire","Data lichidare"])
            randuriProduseDeSters = []
            for i in range(len(df_toateProdusele.index)):
                if int(nrContract) == df_toateProdusele.iloc[i]["Nr. Contract"]:
                    df_produse = df_produse.append(df_toateProdusele.loc[i])
                    randuriProduseDeSters.append(i + 2)
            # Preia date client dupa nr contractului
            df_client = pd.read_excel(listaDB, sheet_name="ListaClienti")
            for i in range(len(df_client.index)):
                print(df_client.iloc[i]["Contracte "], type(df_client.iloc[i]["Contracte "]))
                contract = df_client.iloc[i]["Contracte "]
                if isinstance(contract, int):  # daca clientul are un singur contract in db
                    if df_client.iloc[i]["Contracte "] == int(nrContract):
                        df_client = df_client.loc[i]
                        break
                else:  # clientul are mai multe contracte in db
                    contract = contract.replace(" ", "")
                    contract = contract.split(",")
                    print(contract)
                    if str(nrContract) in contract:
                        df_client = df_client.loc[i]
                        break
            lista = load_workbook(listaDB)
            row = df_contract.values.tolist()
            dataExpirare = datetime.strptime(row[6], '%d.%m.%Y').date()
            if tabelCurent == 0:  # Contractul este in termen
                # adaug Contractul in lista cu contracte incheiate
                row.append("-")
                row.append("-")
                row.append("-")
                row.append("-")
                row.append("-")
                row.append(date.today().strftime('%d.%m.%Y'))
                row.append("CLIENT")
                ws = lista["ListaContracteLichidate"]
                ws.append(row)
                # Mut contractul si in aplicatie
                # Din tabelul cu contracte in termen, in tabelul cu contracte lichidate
                nrRanduri = self.ui.tableWidget_contracteIncheiate.rowCount()
                self.ui.tableWidget_contracteIncheiate.setRowCount(nrRanduri + 1)
                for i in range(len(row)):
                    self.ui.tableWidget_contracteIncheiate.setItem(nrRanduri, i,
                                                                    QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contracteInTermen.removeRow(self.ui.tableWidget_contracteInTermen.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws = lista["ListaContracteInTermen"]
                ws.delete_rows(randContract)
                # Fac aceleasi mutari si pentru produsele din contract
                # Adaug produsele in lista cu produse lichidate de client
                row = df_produse.values.tolist()
                dataExpirare = datetime.strptime(row[0][7], '%d.%m.%Y').date()
                detaliiProduse = []
                numeProdus = []
                obsProdus = []
                sumaImprumutata = []
                for i in range(len(row) - 1, -1, -1):
                    numeProdus.append(row[i][0])
                    obsProdus.append(row[i][1])
                    sumaImprumutata.append(row[i][3])
                    row[i][10]="-"
                    row[i][11]="-"
                    row[i][12]="-"
                    row[i][13]="-"
                    row[i][14]="-"
                    row[i][15]=date.today().strftime('%d.%m.%Y')
                    ws = lista["ListaProduseLichidateClient"]
                    ws.append(row[i])
                    ws = lista["ListaProduseInTermen"]
                    ws.delete_rows(randuriProduseDeSters[i])
                detaliiProduse.append(numeProdus)
                detaliiProduse.append(obsProdus)
                detaliiProduse.append(sumaImprumutata)
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaLichidare(lista,listaDB,nrContract,detaliiProduse)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                    "Contractul " + nrContract + " a fost lichidat!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
            elif tabelCurent == 1:  # Daca contractul este in asteptare
                # adaug Contractul in lista cu contracte prelungite
                row = df_contract.values.tolist()
                row.append("-")
                row.append("-")
                row.append("-")
                row.append("-")
                row.append("-")
                row.append(date.today().strftime('%d.%m.%Y'))
                row.append("CLIENT")
                ws = lista["ListaContracteLichidate"]
                ws.append(row)
                # Mut contractul si in aplicatie
                # Din tabelul cu contracte in asteptare, in tabelul cu contracte lichidate
                nrRanduri = self.ui.tableWidget_contractePrelungite.rowCount()
                self.ui.tableWidget_contracteIncheiate.setRowCount(nrRanduri + 1)
                for i in range(len(row)):
                    self.ui.tableWidget_contracteIncheiate.setItem(nrRanduri, i,
                                                                    QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contracteInAsteptare.removeRow(
                    self.ui.tableWidget_contracteInAsteptare.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws = lista["ListaContracteInAsteptare"]
                ws.delete_rows(randContract)
                # Fac aceleasi mutari si pentru produsele din contract
                # Adaug produsele in lista cu produse lichidate de client
                row = df_produse.values.tolist()
                detaliiProduse = []
                numeProdus = []
                obsProdus = []
                sumaImprumutata = []
                for i in range(len(row) - 1, -1, -1):
                    numeProdus.append(row[i][0])
                    obsProdus.append(row[i][1])
                    sumaImprumutata.append(row[i][3])
                    row[i][10] = "-"
                    row[i][11] = "-"
                    row[i][12] = "-"
                    row[i][13] = "-"
                    row[i][14] = "-"
                    row[i][15] = date.today().strftime('%d.%m.%Y')
                    ws = lista["ListaProduseLichidateClient"]
                    ws.append(row[i])
                    ws = lista["ListaProduseInAsteptare"]
                    ws.delete_rows(randuriProduseDeSters[i])
                detaliiProduse.append(numeProdus)
                detaliiProduse.append(obsProdus)
                detaliiProduse.append(sumaImprumutata)
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaLichidare(lista,listaDB,nrContract,detaliiProduse)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                    "Contractul " + nrContract + " a fost lichidat!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
            elif tabelCurent == 4:  # Daca contractul a mai fost prelungit
                # Fac update in lista de contracte prelungite pentru contractul prelungit
                ws = lista["ListaContracteLichidate"]
                row = df_contract.values.tolist()
                row.append(date.today().strftime('%d.%m.%Y'))
                row.append("CLIENT")
                ws.append(row)
                # Fac update si in tabelul din aplicatie
                # Din tabelul cu contracte prelungite, in tabelul cu contracte prelungite lichidate
                nrRanduri = self.ui.tableWidget_contracteIncheiate.rowCount()
                self.ui.tableWidget_contracteIncheiate.setRowCount(nrRanduri + 1)
                for i in range(len(row)):
                    self.ui.tableWidget_contracteIncheiate.setItem(nrRanduri, i,
                                                                    QtWidgets.QTableWidgetItem(str(row[i])))
                self.ui.tableWidget_contractePrelungite.removeRow(self.ui.tableWidget_contractePrelungite.currentRow())
                # sterg contractul din lista cu contracte in termen
                ws = lista["ListaContractePrelungite"]
                ws.delete_rows(randContract)
                # Fac aceleasi mutari si pentru produsele din contract
                # Mut produsele din lista cu produse prelungite in lista cu produse lichidate client
                row = df_produse.values.tolist()
                detaliiProduse = []
                numeProdus = []
                obsProdus = []
                sumaImprumutata = []
                for i in range(len(row) - 1, -1, -1):
                    numeProdus.append(row[i][0])
                    obsProdus.append(row[i][1])
                    sumaImprumutata.append(row[i][3])
                    row[i][15] = date.today().strftime('%d.%m.%Y')
                    ws = lista["ListaProduseLichidateClient"]
                    ws.append(row[i])
                    ws = lista["ListaProdusePrelungite"]
                    ws.delete_rows(randuriProduseDeSters[i])
                detaliiProduse.append(numeProdus)
                detaliiProduse.append(obsProdus)
                detaliiProduse.append(sumaImprumutata)
                lista.save(listaDB)
                self.PopulareTabelProduse(listaDB)
                self.genereazaNotaLichidare(lista,listaDB,nrContract,detaliiProduse)
                msgBox1 = QMessageBox()
                msgBox1.setIcon(QMessageBox.Information)
                msgBox1.setText(
                    "Contractul " + nrContract + " a fost lichidat!")
                msgBox1.setWindowTitle("Atentie")
                msgBox1.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox1.exec()
                if returnValue == QMessageBox.Ok:
                    self.goHomePage(listaDB)
                    self.ui.lineEdit_numarContractDePrelungit_7.setText("")
                    self.ui.lineEdit_numarContractDePrelungit_6.setText("")
        print("STOP finalizareIncheiereContract- no return:")

    def incheieContractFinal(self,listaDB, nrContract, randContract,indexTabel):
        print("START incheieContractFinal")
        print("se incheie contract incheieContractFinal")
        # Se face update in LISTA CONTRACTE din DB a contractului incheiat ( adaugare "DA" la coloana "INCHEIAT?")
        currentDir = os.getcwd()
        lista = load_workbook(listaDB)
        ws = lista["ListaContracte"]
        ws['F' + str(int(nrContract) + 1)] = "DA"
        # Se face update in LISTA PRODUSE din DB a produselor din contractului incheiat

        # Se face update in tabelul cu contracte INCHEIATE ( contractul incheiat se muta din tabelul cu contracte IN ASTEPTARE in tabelul cu contracte INCHEIATE)
        self.ui.tableWidget_contracteIncheiate.insertRow(0)
        for i in range(self.ui.tableWidget_contracteIncheiate.columnCount()):
            if indexTabel == 1:
                item = self.ui.tableWidget_contracteInAsteptare.item(randContract, i).text()
            else:
                item = self.ui.tableWidget_contracteInTermen.item(randContract, i).text()
            self.ui.tableWidget_contracteIncheiate.setItem(0, i, QtWidgets.QTableWidgetItem(item))
        if indexTabel ==1:
            self.ui.tableWidget_contracteInAsteptare.removeRow(randContract)
        else:
            self.ui.tableWidget_contracteInTermen.removeRow(randContract)
        # Preia produsele din tabelul cu produse in asteptare din aplicatie, pentru a le adauga in tabelul din Nota de Lichidare prin functia genereazaNotaLichidare
        denumireProdus = []
        observatiiProdus = []
        sumaImprumutata = []
        detaliiProduse = []
        self.produseInTermen(listaDB)
        self.produseInAsteptare(listaDB)
        if indexTabel ==1:
            # print("AICI!!!!!!!!!!!!!!!")
            for i in range(self.ui.tableWidget_2.rowCount()):
                # print(i)
                # print("nrContractDINAPLICATIE: ",self.ui.tableWidget_2.item(i, 2).text(),type(self.ui.tableWidget_2.item(i, 2).text()))
                if nrContract == self.ui.tableWidget_2.item(i, 2).text():
                    denumireProdus.append(self.ui.tableWidget_2.item(i, 0).text())
                    observatiiProdus.append(self.ui.tableWidget_2.item(i, 1).text())
                    sumaImprumutata.append(self.ui.tableWidget_2.item(i, 5).text())
        else:
            for i in range(self.ui.tableWidget.rowCount()):
                if nrContract == self.ui.tableWidget.item(i, 2).text():
                    denumireProdus.append(self.ui.tableWidget.item(i, 0).text())
                    observatiiProdus.append(self.ui.tableWidget.item(i, 1).text())
                    sumaImprumutata.append(self.ui.tableWidget.item(i, 5).text())
        detaliiProduse.append(denumireProdus)
        detaliiProduse.append(observatiiProdus)
        detaliiProduse.append(sumaImprumutata)
        # Genereaza si deschide nota de lichidare
        self.genereazaNotaLichidare(lista,listaDB,nrContract,currentDir,detaliiProduse)
        msgBox1 = QMessageBox()
        msgBox1.setIcon(QMessageBox.Information)
        msgBox1.setText(
            "Contractul " + nrContract + " a fost incheiat!")
        msgBox1.setWindowTitle("Atentie")
        msgBox1.setStandardButtons(QMessageBox.Ok)
        returnValue = msgBox1.exec()
        if returnValue == QMessageBox.Ok:
            self.goHomePage(listaDB)
            self.ui.lineEdit_numarContractDePrelungit_7.setText("")
            self.ui.lineEdit_numarContractDePrelungit_6.setText("")
            # self.ui.stackedWidget.setCurrentIndex(0)
        lista.save(currentDir+listaDB)
        print("STOP incheieContractFinal- no return:")
    def genereazaNotaLichidare(self,lista,listaDB,nrContract,detaliiProduse):
        print("START genereazaNotaLichidare- no return:")
        # Genereaza nr nota de lichidare
        numarNotaLichidare = self.genereazaNumarNotaLichidare(lista, listaDB, nrContract)
        lista.save(listaDB)
        currentDir = os.getcwd()
        # 2.Completeaza NOTA DE LICHIDARE
        template = "/templateContract/TemplateNotaDeLichidare.docx"
        copyTemplate = "/noteLichidare"
        originalTempl = currentDir + template
        copyTempl = currentDir + copyTemplate
        shutil.copy(originalTempl, copyTempl)
        document = docx.Document(currentDir +"/noteLichidare/TemplateNotaDeLichidare.docx")
        today = date.today().strftime("%d.%m.%Y")
        numeClient = self.ui.label_numarCercDreapta_4.text() + "-" + self.ui.label_numarCercDreapta_17.text()
        numeClient = numeClient.replace(" ","-")
        numeContract = str(nrContract)+"-"+self.ui.label_numarCercDreapta_4.text()+"-"+self.ui.label_numarCercDreapta_17.text()+".docx"
        cnp = self.ui.label_numarCercDreapta_20.text()
        sumaImprumutata = self.ui.label_numarCercDreapta_12.text()
        comision = self.ui.label_numarCercDreapta_22.text()
        penalizari = self.ui.label_numarCercDreapta_14.text()
        total = self.ui.label_numarCercDreapta_15.text()
        nrBon = self.ui.lineEdit_numarContractDePrelungit_6.text()
        numeNotaLichidare = str(numarNotaLichidare) + "-" + nrContract + "-" + numeClient+".docx"
        numeClient = numeClient.replace("-"," ")
        dataCreareContract = self.ui.label_numarCercDreapta_3.text()
        for paragraph in document.paragraphs:
            runs = paragraph.runs
            for i in range(len(runs)):
                if "[NR_LICHIDARE]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_LICHIDARE]", str(numarNotaLichidare))
                    runs[i].bold = True
                if "[DATA]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[DATA]", today)
                    runs[i].bold = True
                if "[NR_CONTRACT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_CONTRACT]", str(nrContract))
                    runs[i].bold = True
                if "[DATA_CONTRACT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[DATA_CONTRACT]", dataCreareContract)
                    runs[i].bold = True
                if "[NUME_CLIENT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NUME_CLIENT]", numeClient)
                    runs[i].bold = True
                if "[CNP_CLIENT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[CNP_CLIENT]", cnp)
                    runs[i].bold = True
                if "[SUMA_RESTITUIT]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[SUMA_RESTITUIT]", sumaImprumutata)
                    runs[i].bold = True
                if "[COMISION]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[COMISION]", comision)
                    runs[i].bold = True
                if "[PENALIZARI]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[PENALIZARI]", penalizari)
                    runs[i].bold = True
                if "[TOTAL]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[TOTAL]", total)
                    runs[i].bold = True
                if "[NR_BON]" in runs[i].text:
                    runs[i].text = runs[i].text.replace("[NR_BON]", nrBon)
                    runs[i].bold = True
        numarProduse = len(detaliiProduse[0])
        for i in range(2):
            tabel = document.tables[i]
            total = 0
            for j in range(numarProduse):
                cells = tabel.add_row().cells
                cells[0].text = str(j + 1)
                cells[1].text = detaliiProduse[0][j]
                cells[2].text = detaliiProduse[1][j]
                cells[3].text = str(detaliiProduse[2][j])
                total = total + int(detaliiProduse[2][j])
            cells = tabel.rows[1].cells
            cells[3].text = str(total)
            self.moveTotal(tabel)
        # print(currentDir,"currentDir")
        # print(numeNotaLichidare,"numeNotaLichidare")
        document.save(currentDir +"/noteLichidare/" + numeNotaLichidare)
        os.system('open '+currentDir + "/noteLichidare/" + numeNotaLichidare)
        #os.startfile(currentDir + "\\noteLichidare\\" + numeNotaLichidare)
        print("STOP genereazaNotaLichidare- no return:")
    def nrContracteInTermen(self,listaDB):
        print("START nrContracteInTermen")
        listaContracteInTermen = [0, ""]
        listaContracte = ["ListaContracteInTermen","ListaContractePrelungite"]
        for i in range(len(listaContracte)):
            df = pd.read_excel(listaDB, sheet_name=listaContracte[i])
            for i in range(len(df.index)):
                listaContracteInTermen[1] = listaContracteInTermen[1] + str(df.iloc[i]["Nr. Contract"]) + ","
                listaContracteInTermen[0] += 1
        listaContracteInTermen[1] = listaContracteInTermen[1][:-1]
        print("STOP nrContracteInTermen- return:", listaContracteInTermen)
        self.ui.label_numarCercDreapta.setText(str(listaContracteInTermen[0]))
        return listaContracteInTermen
    def nrContracteInAsteptare(self,listaDB):
        print("START nrContracteInAsteptare")
        df = pd.read_excel(listaDB, sheet_name="ListaContracteInAsteptare")
        listaContracteInAsteptare = [0, ""]
        for i in range(len(df.index)):
            listaContracteInAsteptare[1] = listaContracteInAsteptare[1] + str(df.iloc[i]["Nr. Contract"]) + ","
            listaContracteInAsteptare[0] += 1
        listaContracteInAsteptare[1] = listaContracteInAsteptare[1][:-1]
        print("STOP nrContracteInTermen- return:", listaContracteInAsteptare)
        print("STOP nrContracteInAsteptare- return:", listaContracteInAsteptare)
        return listaContracteInAsteptare
    def verificaContracteExpirate(self,listaDB):
        print("START verificaContracteExpirate")
        df = pd.read_excel(listaDB, sheet_name="ListaContracteExpirate")
        listaContracteExpirate = [0, ""]
        for i in range(len(df.index)):
            listaContracteExpirate[1] = listaContracteExpirate[1] + str(df.iloc[i]["Nr. Contract"]) + ","
            listaContracteExpirate[0] += 1
        listaContracteExpirate[1] = listaContracteExpirate[1][:-1]
        print("STOP verificaContracteExpirate- return:", listaContracteExpirate)
        return listaContracteExpirate
    def produseInTermen(self,listaDB):
        print("START produseInTermen")
        listaProduse = ["ListaProduseInTermen","ListaProdusePrelungite"]
        tabel = [self.ui.tableWidget_produseInTermen,self.ui.tableWidget_produsePrelungite]
        listaProduseInTermen = [0, ""]
        for j in range(len(listaProduse)):
            df = pd.read_excel(listaDB, sheet_name=listaProduse[j])
            tabel[j].setRowCount(len(df.index))
            for i in range(len(df.index)):
                tabel[j].setItem(i,0,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Denumire Produs"])))
                tabel[j].setItem(i,1,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Observatii"])))
                tabel[j].setItem(i,2,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Nr. Contract"])))
                tabel[j].setItem(i,3,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data intrare in amanet"])))
                tabel[j].setItem(i,4,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Data expirare"])))
                tabel[j].setItem(i,5,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Suma Imprumutata"])))
                tabel[j].setItem(i,6,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Comision (lei)"])))
                tabel[j].setItem(i,7,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Total de restituit (lei)"])))
                listaProduseInTermen[1] = listaProduseInTermen[1] + str(df.iloc[i]["Denumire Produs"]) + ","
                listaProduseInTermen[0] += 1
        listaProduseInTermen[1] = listaProduseInTermen[1][:-1]
        self.ui.label_numarCercMijloc.setText(str(listaProduseInTermen[0]))
        print("STOP produseInTermen- return:", listaProduseInTermen)
        return listaProduseInTermen
    # def produseInAsteptare(self,listaDB):
    #     print("START produseInAsteptare")
    #     contracteInAsteptare = self.nrContracteInAsteptare(listaDB)
    #     contracteInAsteptare = (contracteInAsteptare[1].split(","))
    #     currentDir = os.getcwd()
    #     lista = load_workbook(currentDir + listaDB)
    #     ws = lista["ListaProduse"]
    #     maxRow = ws.max_row
    #     nrProduse = 0
    #     denumireProdus = []
    #     observatiiProdus = []
    #     nrContract = []
    #     dataIntrare = []
    #     dataExpirare = []
    #     sumaImprumutata = []
    #     comision = []
    #     total = []
    #     for i in range(2, maxRow + 1):
    #         nrContractProdus = ws['K' + str(i)].value
    #         for j in range(len(contracteInAsteptare)):
    #             if not contracteInAsteptare[j] == "":
    #                 if nrContractProdus == int(contracteInAsteptare[j]):
    #                     nrProduse += 1
    #                     numeProdus = ws['B' + str(i)].value
    #                     observatiiProd = ws['C' + str(i)].value
    #                     dataIn = ws['H' + str(i)].value
    #                     dataExp = ws['I' + str(i)].value
    #                     denumireProdus.append(numeProdus)
    #                     observatiiProdus.append(observatiiProd)
    #                     nrContract.append(nrContractProdus)
    #                     dataIntrare.append(dataIn)
    #                     dataExpirare.append(dataExp)
    #                     sumaImprumutata.append(ws['E' + str(i)].value)
    #                     comision.append(ws['G' + str(i)].value)
    #                     total.append(ws['J' + str(i)].value)
    #     self.ui.tableWidget_2.setRowCount(nrProduse)
    #     row = 0
    #     for i in range(len(denumireProdus)):
    #         self.ui.tableWidget_2.setItem( row ,0,QtWidgets.QTableWidgetItem(denumireProdus[i]))
    #         self.ui.tableWidget_2.setItem(row,1, QtWidgets.QTableWidgetItem(str(observatiiProdus[i])))
    #         self.ui.tableWidget_2.setItem(row, 2, QtWidgets.QTableWidgetItem(str(nrContract[i])))
    #         self.ui.tableWidget_2.setItem(row,3, QtWidgets.QTableWidgetItem(str(dataIntrare[i])))
    #         self.ui.tableWidget_2.setItem(row,4, QtWidgets.QTableWidgetItem(str(dataExpirare[i])))
    #         self.ui.tableWidget_2.setItem(row,5, QtWidgets.QTableWidgetItem(str(sumaImprumutata[i])))
    #         self.ui.tableWidget_2.setItem(row,6, QtWidgets.QTableWidgetItem(str(comision[i])))
    #         self.ui.tableWidget_2.setItem(row,7, QtWidgets.QTableWidgetItem(str(total[i])))
    #         row +=1
    #     lista.save(currentDir + listaDB)
    #     print("STOP produseInAsteptare- return:", nrProduse)
    #     return nrProduse

    def nrTotalClienti(self,listaDB):
        print("START nrTotalClienti")
        currentDir = os.getcwd()
        lista = load_workbook(currentDir + listaDB)
        ws = lista["ListaClienti"]
        maxRow = ws.max_row
        print("STOP nrTotalClienti- return:", maxRow-2)
        return maxRow-2
    def nrTotalContracte(self,listaDB):
        print("START cautaClientDupaNrContract")
        currentDir = os.getcwd()
        lista = load_workbook(currentDir + listaDB)
        ws = lista["ListaContracte"]
        maxRow = ws.max_row
        print("STOP nrTotalContracte- return:", maxRow - 2)
        return maxRow - 2
    def preiaDateClient(self):
        print("START preiaDateClient")
        numeClient = self.ui.lineEdit_NumeClient.text().upper()
        prenumeClient = self.ui.lineEdit_PrenumeClient.text().upper()
        serieCi = self.ui.lineEdit_SerieCI.text().upper()
        numarCi = self.ui.lineEdit_NumarCI.text()
        cnp = self.ui.lineEdit_CNP.text()
        eliberatDe = self.ui.lineEdit_EliberatDe.text().upper()
        adresa = self.ui.lineEdit_Adresa.text().upper()
        dataEliberarii = self.ui.dateEdit_dataEliberarii.date().toString("dd.MM.yyyy")
        telefon = self.ui.lineEdit_Telefon.text()
        email = self.ui.lineEdit_Email.text()
        dateClient = [numeClient, prenumeClient, serieCi, numarCi, cnp, eliberatDe, adresa, dataEliberarii, telefon,
                      email]
        print("STOP preiaDateClient- return:", dateClient)
        return dateClient

    def preiaDateProdus(self):
        print("START preiaDateProdus")
        denumireProdus = self.ui.lineEdit_denumireProdus.text()
        observatiiProdus = self.ui.lineEdit_ObservatiiProdus.text()
        valoareEvaluata = self.ui.lineEdit_ValoareEvaluata.text()
        sumaImprumutata = self.ui.lineEdit_SumaImprumutata.text()
        comision = self.ui.lineEdit_ComisionLuna.text()
        dateProdus = [denumireProdus, observatiiProdus, valoareEvaluata, sumaImprumutata, comision]
        print("STOP preiaDateProdus - return:", dateProdus)
        return dateProdus
    # def calculeazaComision(self,sumaImprumutata):
    #     # print("START calculeazaComision")
    #     comisionDorit = 0
    #     sumaImprumutata = int(sumaImprumutata)
    #     if sumaImprumutata > 0 and sumaImprumutata <=99:
    #         comisionDorit = 20
    #     elif sumaImprumutata >=100 and sumaImprumutata <=199:
    #         comisionDorit = 30
    #     elif sumaImprumutata >= 200 and sumaImprumutata <=249:
    #         comisionDorit = 40
    #     elif sumaImprumutata >=250 and sumaImprumutata <=299:
    #         comisionDorit = 50
    #     elif sumaImprumutata >=300 and sumaImprumutata <= 399:
    #         comisionDorit = 60
    #     elif sumaImprumutata >=400 and sumaImprumutata <=499:
    #         comisionDorit = 80
    #     elif sumaImprumutata >= 500 and sumaImprumutata <=599:
    #         comisionDorit = 90
    #     elif sumaImprumutata >= 600 and sumaImprumutata <=699:
    #         comisionDorit = 100
    #     elif sumaImprumutata >= 700 and sumaImprumutata <=799:
    #         comisionDorit = 110
    #     elif sumaImprumutata >= 800 and sumaImprumutata <=899:
    #         comisionDorit = 120
    #     elif sumaImprumutata >= 900 and sumaImprumutata <= 999:
    #         comisionDorit = 130
    #     elif sumaImprumutata >= 1000 and sumaImprumutata <= 1249:
    #         comisionDorit = 140
    #     elif sumaImprumutata >= 1250 and sumaImprumutata <= 1499:
    #         comisionDorit = 160
    #     elif sumaImprumutata >= 1500 and sumaImprumutata <= 1749:
    #         comisionDorit = 190
    #     elif sumaImprumutata >= 1750 and sumaImprumutata <= 1999:
    #         comisionDorit = 220
    #     elif sumaImprumutata >= 2000 and sumaImprumutata <= 2499:
    #         comisionDorit =250
    #     elif sumaImprumutata >= 2500 and sumaImprumutata <= 2999:
    #         comisionDorit = 290
    #     elif sumaImprumutata >= 3000 and sumaImprumutata <= 3499:
    #         comisionDorit = 330
    #     elif sumaImprumutata >= 3500 and sumaImprumutata <= 3999:
    #         comisionDorit = 380
    #     elif sumaImprumutata >= 4000 and sumaImprumutata <= 4499:
    #         comisionDorit = 420
    #     elif sumaImprumutata >= 4500 and sumaImprumutata <= 4999:
    #         comisionDorit = 460
    #     else:
    #         comisionDorit = (10/100)*sumaImprumutata
    #
    #     comisionLeiPerZi = comisionDorit/30
    #     comisionLeiPerZi = round(comisionLeiPerZi,2)
    #     # print("comision dorit:",comisionDorit)
    #     # print("comisionLeiPerZi:",comisionLeiPerZi)
    #     comisionProcentPerZi = (comisionLeiPerZi * 100)/sumaImprumutata
    #     comisionProcentPerZi = round(comisionProcentPerZi,2)
    #     # print("comisionProcentPerZi",comisionProcentPerZi)
    #     # print("STOP calculeazaComision- return:", comisionProcentPerZi)
    #     return comisionProcentPerZi

    def verificareDateClient(self, date):
        print("START verificareDateClient")
        dateNecompletate = 0
        for i in range(len(date)):
            if date[i] == "":
                dateNecompletate = 1
        if dateNecompletate == 1:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu au fost introduse toate datele clientului!")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        if len(date[2]) != 2 or not date[2].isalpha():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Serie C.I incorecta!\nSeria C.I trebuie sa aiba exact 2 caractere alfabetice.\nex: \"NT\"")
            msgBox.setWindowTitle("Atentie la Seria C.I")
            msgBox.setStandardButtons(QMessageBox.Ok)

            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        if len(date[3]) != 6 or not date[3].isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Numar C.I incorect!\nNumar C.I trebuie sa aiba exact 6 cifre.\nex: \"123456\"")
            msgBox.setWindowTitle("Atentie Numar C.I.")
            msgBox.setStandardButtons(QMessageBox.Ok)

            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        if len(date[4]) != 13 or not date[4].isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("CNP incorect!\nCNP-ul trebuie sa aiba exact 13 cifre.\nex: \"1820302270012\"")
            msgBox.setWindowTitle("Atentie la CNP")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        if date[8].find("-") == -1 and not date[8].isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Numarul de telefon incorect!\nTrebuie sa contina doar cifre.\nex: \"0740987654\"")
            msgBox.setWindowTitle("Atentie la numarul de telefon")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        regex = '^[a-z0-9]+[\._]?[a-z0-9]+[@]\w+[.]\w+$'
        if not (re.search(regex, date[9])) and "-" not in date[9]:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Adresa de Email incorecta!.\nex: \"mike.thomson@gmail.com\"")
            msgBox.setWindowTitle("Atentie la Adresa de Email")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateClient- return: false")
            return False
        print("STOP verificareDateClient- return: true")
        return True

    def calculeazaComisionProdus(self, date):
        print("START calculeazaComisionProdus")
        sumaImprumutata = Decimal(date[3])
        sumaImprumutata = round(sumaImprumutata, 2)
        comisionLuna = date[4]
        comisionLeiPeZi = comisionLuna/30
        comisionLeiPeZi = round(comisionLeiPeZi,2)
        comisionProcentePeZi = (comisionLeiPeZi*100)/sumaImprumutata
        comisionProcentePeZi = round(comisionProcentePeZi,2)
        total = sumaImprumutata + comisionLuna
        valoriCalculate = [comisionProcentePeZi, total]
        print("STOP calculeazaComisionProdus - return:", valoriCalculate)
        return valoriCalculate

    def verificareDateProdus(self, date):
        print("START verificareDateProdus")
        dateNecompletate = 0
        for i in range(len(date)):
            if date[i] == "":
                dateNecompletate = 1
        if dateNecompletate == 1:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("Nu au fost introduse toate datele produsului!")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            # print("STOP verificareDateProdus- return: false")
            return False
        if date[2].isdigit() == False:
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Valoarea introdusa pentru \"Valoare evaluata\" trebuie sa contina doar cifre!\nex: \"1000\"")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateProdus- return: false")
            return False
        if not date[3].isdigit():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Valoarea introdusa pentru \"Suma imprumutata\" trebuie sa contina doar cifre!\nex: \"1000\"")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            print("STOP verificareDateProdus- return: false")
            return False
        if not date[4].isdigit() :  #VERIFICARE VALOARE INTRODUSA PENTRU COMISION
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Valoarea introdusa pentru \"Comision\" trebuie sa contina doar cifre!\nex: \"100\"")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            return False
        print("STOP verificareDateProdus- return: true")
        return True

    def styleTabel(self, ws, maxRow, coloana):
        ws[coloana + str(maxRow + 1)].alignment = Alignment(horizontal='center', vertical='center')
        ws[coloana + str(maxRow + 1)].border = Border(left=Side(border_style='thin', color='000000'),
                                                      right=Side(border_style='thin', color='000000'),
                                                      top=Side(border_style='thin', color='000000'),
                                                      bottom=Side(border_style='thin', color='000000'))
    def adaugaClientInDB(self, lista, date, nrContract,listaDB):
        print("START adaugaClientInDB")
        clientulExista = self.verificareClientExistentDupaCNP(listaDB, date)
        ws = lista["ListaClienti"]
        nrRanduri = self.ui.tableWidget_4.rowCount()
        if clientulExista == False:
            row = [date[0],date[1],date[2],int(date[3]),int(date[4]),date[7],date[6],date[8],date[9],nrContract,str(date[5])]
            ws.append(row)
            self.ui.tableWidget_4.setRowCount(nrRanduri+1)
            self.ui.tableWidget_4.setItem(nrRanduri,0,QtWidgets.QTableWidgetItem(date[0]))
            self.ui.tableWidget_4.setItem(nrRanduri,1,QtWidgets.QTableWidgetItem(date[1]))
            self.ui.tableWidget_4.setItem(nrRanduri,2,QtWidgets.QTableWidgetItem(date[4]))
            self.ui.tableWidget_4.setItem(nrRanduri,3,QtWidgets.QTableWidgetItem(date[2]))
            self.ui.tableWidget_4.setItem(nrRanduri,4,QtWidgets.QTableWidgetItem(date[3]))
            self.ui.tableWidget_4.setItem(nrRanduri,5,QtWidgets.QTableWidgetItem(date[6]))
            self.ui.tableWidget_4.setItem(nrRanduri,6,QtWidgets.QTableWidgetItem(date[8]))
            self.ui.tableWidget_4.setItem(nrRanduri,7,QtWidgets.QTableWidgetItem(date[9]))
            self.ui.tableWidget_4.setItem(nrRanduri,8,QtWidgets.QTableWidgetItem(str(nrContract)))
        else:
            contracteExistente = ws['J' + str(clientulExista)].value
            ws['J' + str(clientulExista)] = str(contracteExistente) + " , " + str(nrContract)
            self.ui.tableWidget_4.setItem(clientulExista-2, 8, QtWidgets.QTableWidgetItem( str(contracteExistente) + " , "+str(nrContract)))
            self.styleTabel(ws, clientulExista, 'J')
        print("STOP adaugaClientInDB- no return:")
        lista.save(listaDB)

    def verificareClientExistentDupaCNP(self, listaDB, date):
        print("START verificareClientExistentDupaCNP")
        df = pd.read_excel(listaDB,sheet_name="ListaClienti")
        for i in range(len(df.index)):
            if int(date[4]) == int(df.iloc[i]["CNP"]):
                print("STOP verificareClientExistentDupaCNP- return:", i+2) # i+2 este randul pe care se afla clientul in DB
                return i+2
        return False
    def verificareClientExistentDupaNume(self, lista, date):
        print("START verificareClientExistentDupaNume")
        ws = lista["ListaClienti"]
        maxRow = ws.max_row
        numeClient = date[0]
        prenumeClient = date[1]
        nrClientiCuAcelasiNume = 0
        for i in range(2, maxRow + 2):
            if numeClient == ws['B' + str(i)].value and prenumeClient == ws['C' + str(i)].value:
                nrClientiCuAcelasiNume += 1
        if nrClientiCuAcelasiNume == 0:
            print("STOP verificareClientExistentDupaNume- return: false")
            return False
        elif nrClientiCuAcelasiNume == 1:
            print("STOP verificareClientExistentDupaNume- return: true")
            return True
        else:
            print("STOP verificareClientExistentDupaNume- return: true")
            return True
    def verificareIntroducereNumeClientPentruCautare(self,nume,prenume):
        print("START verificareIntroducereNumeClientPentruCautare")
        if nume == "":
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introduceti numele clientului")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            return False
        elif prenume == "":
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText(
                "Introduceti prenumele clientului")
            msgBox.setWindowTitle("Atentie")
            msgBox.setStandardButtons(QMessageBox.Ok)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                print('OK clicked')
            return False
        print("STOP verificareIntroducereNumeClientPentruCautare- return: true/false")
    def populeazaTabelCautaClient(self,listaDB):
        print("START populeazaTabelCautaClient")
        df = pd.read_excel(listaDB,sheet_name="ListaClienti")
        self.ui.tableWidget_4.setRowCount(len(df.index))
        self.ui.tableWidget_4.setColumnWidth(0,189)
        self.ui.tableWidget_4.setColumnWidth(1,201)
        self.ui.tableWidget_4.setColumnWidth(2,121)
        self.ui.tableWidget_4.setColumnWidth(3,57)
        self.ui.tableWidget_4.setColumnWidth(4,93)
        self.ui.tableWidget_4.setColumnWidth(5,184)
        self.ui.tableWidget_4.setColumnWidth(6,67)
        self.ui.tableWidget_4.setColumnWidth(7,176)
        self.ui.tableWidget_4.setColumnWidth(8,121)
        for i in range(len(df.index)):
            self.ui.tableWidget_4.setItem(i,0,QtWidgets.QTableWidgetItem(df.iloc[i]["Nume"]))
            self.ui.tableWidget_4.setItem(i,1,QtWidgets.QTableWidgetItem(df.iloc[i]["Prenume"]))
            self.ui.tableWidget_4.setItem(i,2,QtWidgets.QTableWidgetItem(str(df.iloc[i]["CNP"])))
            self.ui.tableWidget_4.setItem(i,3,QtWidgets.QTableWidgetItem(df.iloc[i]["Serie C.I"]))
            self.ui.tableWidget_4.setItem(i,4,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Numar C.I"])))
            self.ui.tableWidget_4.setItem(i,5,QtWidgets.QTableWidgetItem(df.iloc[i]["Adresa"]))
            self.ui.tableWidget_4.setItem(i,6,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Telefon"])))
            self.ui.tableWidget_4.setItem(i,7,QtWidgets.QTableWidgetItem(df.iloc[i]["Email"]))
            self.ui.tableWidget_4.setItem(i,8,QtWidgets.QTableWidgetItem(str(df.iloc[i]["Contracte "])))
        print("STOP populeazaTabelCautaClient- no return:")
    def selectClientInTabelDupaCNP(self,valoareCautata):
        print("START selectClientInTabelDupaCNP")
        nrRows = self.ui.tableWidget_4.rowCount()
        for j in range(nrRows):
            if valoareCautata == self.ui.tableWidget_4.item(j, 2).text():
                self.ui.tableWidget_4.selectRow(j)
        print("STOP selectClientInTabelDupaCNP- no return:")
    def selectClientInTabelDupaNume(self,nume,prenume):
        print("START selectClientInTabelDupaNume")
        nrRows = self.ui.tableWidget_4.rowCount()
        for j in range(nrRows):
            if nume == self.ui.tableWidget_4.item(j, 0).text() and prenume == self.ui.tableWidget_4.item(j, 1).text():
                self.ui.tableWidget_4.selectRow(j)
        print("STOP selectClientInTabelDupaNume- no return:")
    def cautaClientInDBDupaCNP(self,CNP,listaDB):
        print("START cautaClientInDBDupaCNP")
        df = pd.read_excel(listaDB,sheet_name="ListaClienti")
        for i in range(len(df.index)):
            if int(CNP) == int(df.iloc[i]["CNP"]):
                self.autocompletCampuri(df, i)
        print("STOP cautaClientInDBDupaCNP - no return:")
    def clickCautaClient(self,listaDB):
        print("START clickCautaClient")
        df = pd.read_excel(listaDB,sheet_name="ListaClienti")
        if self.ui.radioButton.isChecked():
            for i in range(len(df.index)):
                print(df.iloc[i]["CNP"],type(df.iloc[i]["CNP"]))
                if int(self.ui.lineEdit_CautaCnp.text()) == int(df.iloc[i]["CNP"]):
                    self.goToClientGasit()
                    self.selectClientInTabelDupaCNP(self.ui.lineEdit_CautaCnp.text())
                    return True
            self.goToClientNegasit()
            return False
        elif self.ui.radioButton_2.isChecked():
            numeClient = self.ui.lineEdit_CautaNume.text()
            prenumeClient = self.ui.lineEdit_CautaPrenume.text()
            self.verificareIntroducereNumeClientPentruCautare(numeClient,prenumeClient)
            numeClient = numeClient.upper()
            prenumeClient = prenumeClient.upper()
            nrClientiCuAcelasiNume = 0
            linieClient = 0
            for i in range(len(df.index)):
                if numeClient == df.iloc[i]["Nume"] and prenumeClient == df.iloc[i]["Prenume"]:
                    nrClientiCuAcelasiNume += 1
                    linieClient = i
            if nrClientiCuAcelasiNume == 0:
                self.goToClientNegasit()
                return False
            elif nrClientiCuAcelasiNume == 1:
                self.goToClientGasit()
                self.selectClientInTabelDupaNume(numeClient,prenumeClient)
                return True
            else:
                self.ui.label_numarCercMijloc_4.setText(str(nrClientiCuAcelasiNume))
                self.goToMaiMultiClienti()
                return False
        print("STOP clickCautaClient- return: true/false")

    def autocompletCampuri(self,df,i):
        print("START autocompletCampuri")
        self.ui.lineEdit_NumeClient.setText(df.iloc[i]["Nume"])
        self.ui.lineEdit_PrenumeClient.setText(df.iloc[i]["Prenume"])
        self.ui.lineEdit_SerieCI.setText(df.iloc[i]["Serie C.I"])
        self.ui.lineEdit_NumarCI.setText(str(df.iloc[i]["Numar C.I"]))
        self.ui.lineEdit_CNP.setText(str(df.iloc[i]["CNP"]))
        self.ui.lineEdit_EliberatDe.setText(df.iloc[i]["C.I Eliberat de"])
        data = datetime.strptime(str(df.iloc[i]["Data eliberare C.I"]),'%d.%m.%Y')
        self.ui.dateEdit_dataEliberarii.setDate(data)
        self.ui.lineEdit_Adresa.setText(df.iloc[i]["Adresa"])
        self.ui.lineEdit_Telefon.setText(str(df.iloc[i]["Telefon"]))
        self.ui.lineEdit_Email.setText(df.iloc[i]["Email"])
        print("STOP autocompletCampuri- no return:")
    def genereazaNumarNotaLichidare(self,lista,listaDB,nrContract):
        print("START genereazaNumarContract")
        ws = lista["ListaNoteLichidare"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        ws['C' + str(maxRow + 1)] = date.today().strftime("%d.%m.%Y")
        ws['B' + str(maxRow + 1)] = nrContract
        self.styleTabel(ws, maxRow, 'A')
        self.styleTabel(ws, maxRow, 'B')
        self.styleTabel(ws, maxRow, 'C')
        lista.save(listaDB)
        print("STOP genereazaNumarContract- return:", maxRow)
        return maxRow
    def genereazaNumarContract(self, lista):
        print("START genereazaNumarContract")
        ws = lista["ListaNrContracte"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        self.styleTabel(ws, maxRow, 'A')
        print("STOP genereazaNumarContract- return:", maxRow)
        return maxRow
    def genereazaNumarDispoziteDePlata(self, lista,nrContract):
        print("START genereazaNumarDispoziteDePlata")
        ws = lista["DispozitiiPlata"]
        maxRow = ws.max_row
        ws['A' + str(maxRow + 1)] = maxRow
        ws['B' + str(maxRow + 1)] = date.today().strftime("%d.%m.%Y")
        ws['C' + str(maxRow + 1)] = nrContract
        self.styleTabel(ws, maxRow, 'A')
        self.styleTabel(ws, maxRow, 'B')
        self.styleTabel(ws, maxRow, 'C')
        print("STOP genereazaNumarDispoziteDePlata- return:", maxRow)
        return maxRow
    def nrUltimulContract(self, lista,listaDB):
        print("START nrUltimulContract")
        ws = lista["ListaNrContracte"]
        maxRow = ws.max_row
        lista.save(listaDB)
        print("STOP nrUltimulContract- return:", maxRow -1)
        return maxRow -1

    def adaugaDateClientInContract(self,listaDB):
        print("START adaugaDateClientInContract")
        currentDir = os.getcwd();
        if self.ui.radioButton_4.isChecked():
            document = docx.Document(currentDir + "/Contracte/TemplateContractAmanet.docx")
        else:
            document = docx.Document(currentDir + "/Contracte/TemplateContractVanzare.docx")
        lista = load_workbook(listaDB)
        nrContract = self.genereazaNumarContract(lista)
        if self.ui.radioButton_4.isChecked():
            nrDispPlata = self.genereazaNumarDispoziteDePlata(lista,nrContract)
        else:
            nrDispPlata = "0"
        dateClient = self.preiaDateClient()
        if self.verificareDateClient(dateClient):
            today = date.today().strftime("%d.%m.%Y")
            if self.ui.radioButton_4.isChecked():
                endDay = (datetime.now() + timedelta(days=30)).strftime('%d.%m.%Y')
            else:
                endDay = (datetime.now() + timedelta(days=-6)).strftime('%d.%m.%Y')
            for paragraph in document.paragraphs:
                runs = paragraph.runs
                for i in range(len(runs)):
                    if "[NUMAR_CONTRACT]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[NUMAR_CONTRACT]", str(nrContract))
                        runs[i].bold = True
                    if "[DATA]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[DATA]", today)
                        runs[i].bold = True
                    if "[NUME]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[NUME]", dateClient[0])
                        runs[i].bold = True
                    if "[PRENUME]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[PRENUME]", dateClient[1])
                        runs[i].bold = True
                    if "[CNP]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[CNP]", dateClient[4])
                        runs[i].bold = True
                    if "[SERIE_CI]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[SERIE_CI]", dateClient[2])
                        runs[i].bold = True
                    if "[NUMAR_CI]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[NUMAR_CI]", dateClient[3])
                        runs[i].bold = True
                    if "[ELIBERAT_DE]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[ELIBERAT_DE]", dateClient[5])
                        runs[i].bold = True
                    if "[DATA_ELIBERARII]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[DATA_ELIBERARII]", dateClient[7])
                        runs[i].bold = True
                    if "[ADRESA]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[ADRESA]", dateClient[6])
                        runs[i].bold = True
                    if "[DATA_SCADENTA]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[DATA_SCADENTA]", endDay)
                        runs[i].bold = True
                    if "[TELEFON]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[TELEFON]", dateClient[8])
                        runs[i].bold = True
                    if "[EMAIL]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[EMAIL]", dateClient[9])
                        runs[i].bold = True
                    if "[NR_DISPOZITIE_PLATA]" in runs[i].text:
                        runs[i].text = runs[i].text.replace("[NR_DISPOZITIE_PLATA]", str(nrDispPlata))
                        runs[i].bold = True
            contractName = str(nrContract) + "-" + dateClient[0] + "-" + dateClient[1] + ".docx"
            contractName = contractName.replace(" ","-")
            self.adaugaClientInDB(lista, dateClient, nrContract,listaDB)
            document.save(currentDir + "/Contracte/" + contractName)
            lista.save(listaDB)
            self.ui.stackedWidget.setCurrentIndex(2)
            return True
        else:
            return False
        print("STOP adaugaDateClientInContract - return: true/false")

    def remove_row(self, tabel, row):
        tbl = tabel._tbl
        tr = row._tr
        tbl.remove(tr)
        return tr

    def moveTotal(self, tabel):
        row = tabel.rows[1]
        tr = self.remove_row(tabel, row)
        tabel.rows[len(tabel.rows) - 1]._tr.addnext(tr)
    def adaugaProdusInDB(self,lista,dateProd,nrUltimulContract,listaDB,valoriCalculate):
            print("START adaugaProdusInDB")
            if self.ui.radioButton_4.isChecked():
                ws = lista["ListaProduseInTermen"]
                row = [dateProd[0], dateProd[1], dateProd[2], dateProd[3], valoriCalculate[0], dateProd[4],
                       date.today().strftime("%d.%m.%Y")
                    , (date.today() + timedelta(days=30)).strftime('%d.%m.%Y'), valoriCalculate[1], nrUltimulContract]
                ws.append(row)
            else:
                ws = lista["ListaProduseExpirate"]
                row = [dateProd[0], dateProd[1], dateProd[2], dateProd[3], valoriCalculate[0], dateProd[4],
                       date.today().strftime("%d.%m.%Y")
                    , (date.today() + timedelta(days=-6)).strftime('%d.%m.%Y'), valoriCalculate[1], nrUltimulContract]
                ws.append(row)
            lista.save(listaDB)
            self.PopulareTabelProduse(listaDB)
            print("STOP adaugaProdusInDB - no return:")
    def adaugaDateProdusInContract(self,listaDB):
            print("START adaugaDateProdusInContract")
            currentDir = os.getcwd();
            dateClient = self.preiaDateClient()
            lista = load_workbook(listaDB)
            nrUltimulContract = self.nrUltimulContract(lista,listaDB)
            if self.verificareDateClient(dateClient):
                    currentDir = os.getcwd();
                    contractName = str(nrUltimulContract) + "-" + dateClient[0] + "-" + dateClient[1] + ".docx"
                    contractName = contractName.replace(" ","-")
                    document = docx.Document(currentDir + "/contracte/" + contractName)
                    tabel = document.tables[0]
                    dateProdus = self.preiaDateProdus()
                    if self.verificareDateProdus(dateProdus):
                            dateProdus[2] = round(Decimal(dateProdus[2]), 2)
                            dateProdus[3] = round(Decimal(dateProdus[3]), 2)
                            dateProdus[4] = round(Decimal(dateProdus[4]), 2)
                            nrProdus = len(tabel.rows) - 1
                            valoriCalculate = self.calculeazaComisionProdus(dateProdus)
                            cells = tabel.add_row().cells
                            cells[0].text = str(nrProdus)
                            cells[1].text = dateProdus[0]
                            cells[2].text = dateProdus[1]
                            cells[3].text = str(valoriCalculate[0])
                            cells[4].text = str(dateProdus[4])
                            cells[5].text = str(dateProdus[2])
                            cells[6].text = str(dateProdus[3])
                            cells[7].text = str(valoriCalculate[1])
                            self.adaugaProdusInDB(lista,dateProdus,nrUltimulContract,listaDB,valoriCalculate)
                            self.ui.lineEdit_denumireProdus.setText("")
                            self.ui.lineEdit_ObservatiiProdus.setText("")
                            self.ui.lineEdit_ValoareEvaluata.setText("")
                            self.ui.lineEdit_SumaImprumutata.setText("")
                            self.ui.lineEdit_ComisionLuna.setText("")
                            document.save(currentDir + "/Contracte/" + contractName)
                            return True
                    else:
                            return False
            print("STOP adaugaDateProdusInContract - return: true/false")
            lista.save(listaDB)

    def Total(self, tabel):
            print("START Total")
            sumComisionLei = 0
            sumValEval = 0
            sumImprumut = 0
            sumTotal = 0
            # print("len tabel",len(tabel.rows))
            for i in range(1, len(tabel.rows) - 1):
                # print(i)
                sumComisionLei = sumComisionLei + round(Decimal(tabel.cell(i, 4).text),2)
                sumValEval = sumValEval + Decimal(tabel.cell(i, 5).text)
                sumImprumut = sumImprumut + Decimal(tabel.cell(i, 6).text)
                sumTotal = sumTotal + round(Decimal(tabel.cell(i, 7).text),2)
                # print("SUMA TOTALA: ",sumTotal)
            # print(sumComisionLei,sumValEval,sumImprumut,sumTotal)
            valoriTotale = [sumComisionLei, sumValEval, sumImprumut, sumTotal]
            # print("####valoritotale",valoriTotale)
            comisionTotalProcentPerZi = valoriTotale[0]/30
            comisionTotalProcentPerZi = round(comisionTotalProcentPerZi,2)
            comisionTotalProcentPerZi = (comisionTotalProcentPerZi *100)/valoriTotale[2]
            comisionTotalProcentPerZi = round(comisionTotalProcentPerZi,2)
            #comisionTotalProcentPerZi = ((valoriTotale[0] *100)/valoriTotale[2])/30
            # print(comisionTotalProcentPerZi)
            comisionTotalProcentPerZi = round(comisionTotalProcentPerZi,2)
            # print(comisionTotalProcentPerZi)
            # print(comisionTotalProcentPerZi)
            valoriTotale = [sumComisionLei, sumValEval, sumImprumut, sumTotal,comisionTotalProcentPerZi]
            print("STOP TOTAL- return:", valoriTotale)
            return valoriTotale
    def updateTabelContracte(self,nrContract,nume,prenume,telefon,email,sumaImprumutata,comision,total):
        print("START updateTabelContracte")
        dataAzi = date.today()
        dataAzi = dataAzi.strftime("%d.%m.%Y")
        dataPeste30Zile = (date.today() + timedelta(days=30))
        dataPeste30Zile = dataPeste30Zile.strftime("%d.%m.%Y")
        nrRanduriInainte  = self.ui.tableWidget_contracteInTermen.rowCount()
        self.ui.tableWidget_contracteInTermen.setRowCount(nrRanduriInainte+1)
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 0, QtWidgets.QTableWidgetItem(str(nrContract)))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 1, QtWidgets.QTableWidgetItem(nume))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 2, QtWidgets.QTableWidgetItem(prenume))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 3, QtWidgets.QTableWidgetItem(str(telefon)))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 4, QtWidgets.QTableWidgetItem(email))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 5, QtWidgets.QTableWidgetItem(dataAzi))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 6, QtWidgets.QTableWidgetItem(dataPeste30Zile))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 7, QtWidgets.QTableWidgetItem(sumaImprumutata))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 8, QtWidgets.QTableWidgetItem(comision))
        self.ui.tableWidget_contracteInTermen.setItem(nrRanduriInainte, 9, QtWidgets.QTableWidgetItem(total))
        print("STOP updateTabelContracte - no return:")
    def Finalizare(self,listaDB):
            print("START Finalizare")
            if self.adaugaDateProdusInContract(listaDB):
                    dateClient = self.preiaDateClient()
                    currentDir = os.getcwd();
                    lista = load_workbook(listaDB)
                    nrUltimulContract = self.nrUltimulContract(lista,listaDB)
                    contractName =str(nrUltimulContract) + "-" + dateClient[0] + "-" + dateClient[1] + ".docx"
                    contractName = contractName.replace(" ","-")
                    document = docx.Document(currentDir + "/Contracte/" + contractName)
                    tabel = document.tables[0]
                    self.moveTotal(tabel)
                    valoriTotale = self.Total(tabel)
                    tabel.cell(len(tabel.rows) - 1, 4).text = str(valoriTotale[0])
                    tabel.cell(len(tabel.rows) - 1, 5).text = str(valoriTotale[1])
                    tabel.cell(len(tabel.rows) - 1, 6).text = str(valoriTotale[2])
                    tabel.cell(len(tabel.rows) - 1, 7).text = str(valoriTotale[3])
                    comision10zile = (valoriTotale[0]/30)*10
                    comision10zile = round(comision10zile,2)
                    comisionLeiPeZi = valoriTotale[0]/30
                    comisionLeiPeZi = round(comisionLeiPeZi,2)
                    if self.ui.radioButton_4.isChecked():
                        endDay = (date.today() + timedelta(days=30)).strftime('%d.%m.%Y')
                    else:
                        endDay = (date.today() + timedelta(days=-6)).strftime('%d.%m.%Y')
                    today = date.today().strftime('%d.%m.%Y')
                    if self.ui.radioButton_4.isChecked():
                        ws = lista["ListaContracteInTermen"]
                    else:
                        ws = lista["ListaContracteExpirate"]
                    row = [nrUltimulContract, dateClient[0], dateClient[1], dateClient[8], dateClient[9], today, endDay,valoriTotale[2],valoriTotale[0],valoriTotale[3]]
                    ws.append(row)
                    lista.save(listaDB)
                    # self.updateTabelContracte(nrUltimulContract,dateClient[0],dateClient[1],dateClient[8],dateClient[9],str(valoriTotale[2]),str(valoriTotale[0]), str(valoriTotale[3]))
                    self.PopulareTabelContracte(listaDB)
                    for paragraph in document.paragraphs:
                        runs = paragraph.runs
                        for i in range(len(runs)):
                            if "[SUMA_IMPRUMUTATA]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[SUMA_IMPRUMUTATA]",
                                                                    str(valoriTotale[2]))
                                runs[i].bold = True
                            if "[COMISION]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[COMISION]", str(valoriTotale[0]))
                                runs[i].bold = True
                            if "[TOTAL]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[TOTAL]",
                                                                    str(valoriTotale[3]))
                                runs[i].bold = True
                            if "[COMISION_ZI]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[COMISION_ZI]",
                                                                    str(valoriTotale[4]))
                                runs[i].bold = True
                            if "[comision_10_zile]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[comision_10_zile]",
                                                                    str(comision10zile))
                                runs[i].bold = True
                            if "[COMISION_LEI_ZI]" in runs[i].text:
                                runs[i].text = runs[i].text.replace("[COMISION_LEI_ZI]",str(comisionLeiPeZi))
                                runs[i].bold = True
                    document.save(currentDir + "/Contracte/" + contractName)
                    lista.save(listaDB)
                    self.ui.lineEdit_NumeClient.setText("")
                    self.ui.lineEdit_PrenumeClient.setText("")
                    self.ui.lineEdit_NumarCI.setText("")
                    self.ui.lineEdit_SerieCI.setText("")
                    self.ui.lineEdit_CNP.setText("")
                    self.ui.lineEdit_EliberatDe.setText("")
                    self.ui.lineEdit_Adresa.setText("")
                    self.ui.lineEdit_Telefon.setText("")
                    self.ui.lineEdit_Email.setText("")
                    self.ui.lineEdit_ComisionLuna.setText("")
                    msgBox = QMessageBox()
                    msgBox.setIcon(QMessageBox.Information)
                    msgBox.setText("Contractul a fost generat cu succes.")
                    msgBox.setWindowTitle("Atentie")
                    msgBox.setStandardButtons(QMessageBox.Ok)
                    returnValue = msgBox.exec()
                    if returnValue == QMessageBox.Ok:
                            print('OK clicked')
                    # print(currentDir)
                    os.system('open '+currentDir + "/contracte/" + contractName)
                    #os.startfile(currentDir+"/contracte/"+contractName,"open")
                    self.goHomePage(listaDB)
                    # os.startfile(currentDir + "\\contracte\\" + contractName)
            print("STOP FINALIZARE - no return")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec_())
